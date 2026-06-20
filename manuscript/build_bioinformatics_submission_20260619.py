from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "bioinformatics_application_note_submission_20260619"
FIG_DIR = OUT / "figures"
TABLE_DIR = OUT / "tables"

MAIN_MD = """# cfRNA-BrainTrace: hierarchical brain-origin inference from RNA-seq with a primate transcriptomic atlas

**Article type:** Application Note  
**Category:** Gene expression  
**Authors and affiliations:** Author information to be supplied in the submission system  
**Corresponding author:** Corresponding author details to be supplied in the submission system

## Abstract

### Summary

cfRNA-BrainTrace is a Python and Streamlit application for hierarchical brain-origin candidate ranking from RNA expression profiles. Its formal route uses projected-VSD expression only to generate a broad Network Top3 beam, then returns to logCPM-compatible expression for resolution-group and exact-region reranking. The software reports anatomical resolution, confidence and scope diagnostics so that stable coarse signals are not overstated as deterministic exact-region calls.

### Availability and Implementation

Implemented in Python 3.11+ with command-line and Streamlit interfaces. Source code is hosted in the development repository **https://github.com/wz7717/cfrna-brain-tracing**. A versioned public release, archival DOI, licence statement and web service URL will accompany the final submission package.

### Contact

Corresponding author contact details will be provided in the submission system.

### Supplementary information

Supplementary data are available online.

## Text

### Motivation

RNA expression profiles can retain tissue-of-origin information, but within-brain source tracing is limited by transcriptional similarity among neighbouring regions, atlas granularity and domain shifts between reference tissue, tumour tissue and biofluid RNA. These constraints are especially important for atlas-based brain RNA tracing: a stable broad candidate may be recoverable even when exact-region localization is not. A single exact-region output can therefore imply more anatomical resolution than the data support. cfRNA-BrainTrace addresses this problem by returning a hierarchy of Network, resolution-group and exact-region candidates with explicit confidence, coverage and scope diagnostics.

### System and methods

The formal route was selected after comparing three behaviours observed during validation. First, projected-VSD query representation was consistently useful for broad Network candidate generation. Second, direct exact-region scoring in projected-VSD space was not consistently superior to native or logCPM-based alternatives. Third, local reranking within the Network beam improved interpretability at downstream anatomical resolutions. The resulting locked route is therefore a hybrid rather than a single-scale classifier.

Operationally, cfRNA-BrainTrace first normalizes the uploaded expression table into logCPM- or logTPM-compatible space and aligns genes to the reference panel. The query is then projected into Bo2023-like variance-stabilized space for 10-class Network scoring only. The top three Networks form a candidate beam; candidate regions outside this beam are excluded from downstream regional ranking. Within the retained beam, resolution groups are reranked using local logCPM-compatible expression, and exact regions are then reranked as lower-confidence local candidates. This route uses the Bo2023 macaque brain transcriptomic reference without converting the atlas into a new projected atlas. It is implemented in `core/network_tracing.py` and `core/bo2023_region_tracing.py`, with the same scoring core used by the command-line and Streamlit interfaces.

### Validation and outputs

Validation was organized around the same route logic used by the software. Network-level experiments first tested whether projected-VSD queries were suitable for generating the broad candidate beam. In fold-local leave-one-sample-out Network validation, projected VSD achieved Top1/Top3 accuracy of 58.00%/91.58%, exceeding logCPM baseline and native VSD routes. In strict leave-one-monkey-out validation, projected VSD reached 53.72%/91.33%. In contrast, direct exact-region projected-VSD scoring was not consistently superior, especially in leave-one-monkey-out evaluation, so exact-region scoring was not used as the sole endpoint.

The complete validation then tested the formal three-tier procedure end to end: projected-VSD Network Top3 beam generation, logCPM-compatible resolution-group reranking and logCPM-compatible exact-region reranking. This route achieved Network Top3 accuracy of 92.38% in leave-one-sample-out validation and 91.21% in leave-one-monkey-out validation. Resolution-group Top3 accuracy was 72.36% and 69.09%, respectively, whereas exact-region Top3 was 45.33% and 42.36%. The accuracy gradient across levels supports Network Top3 as the primary endpoint and resolution group as the most defensible region-level output. Exact-region results are retained as exploratory local candidate rankings rather than deterministic localization.

External analyses were matched to the resolution supported by their labels. In Allen Human Brain Atlas mapped-label validation, the hybrid route achieved Network Top1/Top3 accuracy of 74.68%/94.42%, resolution-group Top1/Top3 accuracy of 36.26%/67.03% and exact Top1/Top3 accuracy of 24.18%/42.86%; hybrid exact Top3 exceeded both logCPM baseline and projected-VSD-only scoring, supporting the use of projected VSD for the broad beam and logCPM expression for local reranking. In TCGA/BraTS glioma tissue RNA-seq with MRI-derived labels, results support only coarse anatomical consistency because MRI truth labels are human atlas labels rather than Bo2023 macaque exact-region identifiers. In that setting, hybrid Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%. GSE189919 was used to verify projection feasibility rather than accuracy because anatomical truth labels were unavailable.

### Use and limitations

The software exports ranked candidates, matched-gene coverage, entropy, margin, route identifiers and resolution-specific warnings. Accuracy is calculated only when independent anatomical truth is available. Unlabelled biofluid predictions are reported as transfer stress tests or hypothesis-generating candidate rankings rather than clinical localization. The final release package is planned to include non-commercial access, test data, a stable public URL, versioned source code and an archival DOI.

## Funding

Funding information will be declared in the submission system.

## Conflict of Interest

None declared.

## Data availability

The Bo2023 macaque transcriptomic atlas, Allen Human Brain Atlas, TCGA/BraTS and GEO datasets are available from their original repositories under their original access conditions. Processed non-sensitive evaluation tables and figure source data are prepared for inclusion in the tagged public release and archive.

## References

Bakas,S. *et al.* (2017) Advancing The Cancer Genome Atlas glioma MRI collections with expert segmentation labels and radiomic features. *Sci. Data*, **4**, 170117.

Bo,T. *et al.* (2023) Brain-wide and cell-specific transcriptomic insights into MRI-derived cortical morphology in macaque monkeys. *Nat. Commun.*, **14**, 1283.

Hawrylycz,M.J. *et al.* (2012) An anatomically comprehensive atlas of the adult human brain transcriptome. *Nature*, **489**, 391-399.

Vorperian,S.K. *et al.* (2022) Cell types of origin of the cell-free transcriptome. *Nat. Biotechnol.*, **40**, 855-861.

## Figure legend

**Figure 1. cfRNA-BrainTrace three-tier route and validation evidence.** The query is projected into Bo2023-like VSD space only for Network Top3 beam generation; downstream resolution-group and exact-region reranking is performed in logCPM-compatible local expression space. Internal validation supports high Network Top3 accuracy in both LOSO and LOMO settings, while resolution-group accuracy is consistently higher than exact-region accuracy. AHBA supports mapped-label external validation, whereas TCGA/BraTS supports only coarse anatomical consistency.

**Alt text:** Multi-panel figure summarizing the cfRNA-BrainTrace workflow and validation. A workflow diagram shows query preprocessing, projected-VSD Network Top3 beam generation and logCPM local reranking. Bar charts show Network Top3 near 92% in internal validation, resolution-group Top3 around 69%-72% and lower exact-region Top3 around 42%-45%. External panels show AHBA Network Top3 of 94.42% and TCGA/BraTS broad-anatomy Top3 of 64.62%.
"""

SUPP_MD = """# Supplementary Material for cfRNA-BrainTrace

## Supplementary Methods

### S1 Formal validation route

All validation analyses followed the same formal route used by the software. A query profile was first represented in logCPM- or logTPM-compatible expression space and then projected into Bo2023-like VSD space for Network scoring only. The top three Networks formed the candidate beam. Resolution-group and exact-region candidates were then reranked within that beam using local logCPM-compatible expression. This design separates broad candidate generation from fine regional interpretation.

### S2 Internal validation design

Internal validation used the Bo2023 macaque brain reference. Two settings were evaluated. Leave-one-sample-out validation tested whether the route could recover the held-out sample label when other samples from the reference remained available. Leave-one-monkey-out validation held out all samples from one animal at a time and therefore tested donor-level generalization. Metrics were reported separately for Network, resolution-group and exact-region levels. Top1, Top3 and median true-rank were used to distinguish exact calls from candidate-list recovery.

### S3 External validation design

External validation was limited by the label resolution available in each dataset. AHBA human brain RNA-seq was used for mapped-label validation because its anatomical labels could be harmonized to Network, resolution-group and a subset of exact-region labels. TCGA/BraTS glioma tissue RNA-seq with MRI-derived labels was used only for coarse anatomical consistency because its truth labels are human imaging labels, not Bo2023 macaque exact-region identifiers. GSE189919 was used to test whether an external matrix could be projected into the model gene space; it was not used for accuracy estimation because patient-level anatomical truth was unavailable.

### S4 Figures and tables

Supplementary Figure S1 provides the main route and validation summary artwork. Supplementary Tables S1-S5 provide the internal validation design, internal validation results, external validation design, external validation results and figure/table index. Together, these materials document both how the route was tested and what result each validation setting supports.

## Supplementary Results

### S1 Internal route selection

The first internal analysis evaluated Network-level candidate generation. Projected VSD achieved Network Top1/Top3 of 58.00%/91.58% in LOSO and 53.72%/91.33% in LOMO, exceeding logCPM baseline and native VSD at Network Top3. Direct exact-region scoring was lower and less stable, especially in LOMO. These results motivated the use of projected VSD for broad Network beam generation and logCPM-compatible expression for downstream local reranking.

### S2 Formal internal three-tier validation

In the complete LOSO validation, the formal route achieved Network, resolution-group and exact-region Top3 values of 92.38%, 72.36% and 45.33%. In complete LOMO validation, the corresponding Top3 values were 91.21%, 69.09% and 42.36%. Median true-rank increased from Network to exact-region levels, consistent with decreasing anatomical certainty at finer resolution. Resolution group is therefore the preferred region-level endpoint, while exact-region output is retained as a candidate ranking.

### S3 External validation results

In AHBA, the hybrid route achieved Network Top1/Top3 of 74.68%/94.42%, resolution-group Top1/Top3 of 36.26%/67.03% and exact-region Top1/Top3 of 24.18%/42.86%. Hybrid exact Top3 exceeded logCPM baseline (30.77%) and projected-VSD-only scoring (29.67%). In TCGA/BraTS, hybrid Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%, supporting only coarse anatomical consistency. GSE189919 overlapped 15,622/21,668 projector genes, corresponding to 72.10% gene-space coverage, supporting projection feasibility rather than source-localization accuracy.
"""

CHECKLIST_MD = """# Bioinformatics Application Note submission checklist

Source: Oxford Academic Bioinformatics Author Guidelines, checked 2026-06-19.

## Implemented in this draft

- Application Note format with title page, short structured abstract and text.
- Four abstract headings: Summary; Availability and Implementation; Contact; Supplementary information.
- Main manuscript kept to one figure and no main table.
- Figure 1 embedded in the manuscript and exported separately.
- Figure 1 width set for double-column presentation (178 mm / 7.0 in).
- Figure has white background, adjacent legend and `Alt text:` immediately below the legend.
- Supplementary figures, tables and methods are consolidated into one PDF.
- Latest 2026-06-19 validation route is used: projected-VSD Network Top3 beam followed by logCPM resolution-group and exact-region reranking.

## Blocking items before real submission

- Replace author, affiliation, contact and funding placeholders.
- Make the GitHub repository public or otherwise freely available to non-commercial users at a stable URL.
- Add tagged release, test data and archival DOI, for example Zenodo or Figshare.
- Add final web-service URL if submitting as an accessible web interface.
- Confirm OSI-approved licence and maintenance commitment.
- Decide whether AI/tool assistance must be disclosed in cover letter, acknowledgements or supplementary material under journal policy.
"""

CN_MD = """# cfRNA-BrainTrace：基于灵长类转录组图谱的 RNA-seq 分层脑来源推断

**文章类型：** Application Note  
**栏目：** Gene expression  
**作者与单位：** 作者信息将在投稿系统中提供  
**通讯作者：** 通讯作者信息将在投稿系统中提供

## 摘要

### Summary

cfRNA-BrainTrace 是一个 Python 与 Streamlit 应用，用于从 RNA 表达谱中进行分层脑来源候选排序。正式路线只在 Network Top3 beam 生成阶段使用 projected-VSD 表达，随后回到 logCPM 兼容表达空间进行 resolution-group 和 exact-region 重排序。软件同时报告解剖分辨率、置信度和适用边界，避免把稳定的粗粒度信号误写成确定性 exact-region 定位。

### Availability and Implementation

软件基于 Python 3.11+ 实现，支持命令行和 Streamlit 网页界面。源码托管于开发仓库 **https://github.com/wz7717/cfrna-brain-tracing**。最终投稿包将配套公开版本、归档 DOI、许可证声明和在线服务地址。

### Contact

通讯作者联系方式将在投稿系统中提供。

### Supplementary information

补充材料在线提供。

## 正文

### 研究动机

RNA 表达谱可能保留组织来源信息，但脑内溯源受到相邻脑区转录相似性、图谱粒度以及组织、肿瘤和 biofluid 样本域偏移的限制。对于 atlas-based brain RNA tracing，一个粗粒度候选可能较稳定，但 exact-region localization 未必成立。如果只输出单一 exact-region 标签，容易暗示数据并不支持的解剖分辨率。cfRNA-BrainTrace 因此返回 Network、resolution group 和 exact region 三个层级的候选结果，并同步给出置信度、覆盖度和适用范围诊断。

### 系统与方法

正式路线来自三组验证现象。第一，projected-VSD query representation 对 broad Network candidate generation 最稳定。第二，直接在 projected-VSD 空间做 exact-region scoring 并不总是优于 native VSD 或 logCPM-based alternatives。第三，在 Network beam 内进行 local reranking 能提高下游解剖层级的解释性。因此最终路线是一个 hybrid procedure，而不是单一表达尺度上的 classifier。

在实际操作中，cfRNA-BrainTrace 先将上传表达表标准化为 logCPM/logTPM 兼容空间，并与参考基因面板对齐。随后 query 被投影到 Bo2023-like VSD 空间，仅用于 10-class Network scoring。Top3 Networks 构成 candidate beam，beam 之外的候选脑区不进入下游 regional ranking。在保留的 beam 内，resolution group 使用 logCPM 兼容的局部表达空间重排序，exact region 再作为较低置信度的局部候选排序输出。该路线使用 Bo2023 猕猴脑转录组参考，但不把整个 atlas 转换成新的 projected atlas。主线实现位于 `core/network_tracing.py` 和 `core/bo2023_region_tracing.py`，命令行和 Streamlit 调用同一评分核心。

### 验证与输出

验证设计与软件路线保持一致。首先在 Network 层测试 projected-VSD query 是否适合作为 broad candidate beam。在 fold-local leave-one-sample-out Network 验证中，projected VSD 的 Top1/Top3 为 58.00%/91.58%，高于 logCPM baseline 和 native VSD。在 strict leave-one-monkey-out 验证中，projected VSD 达到 53.72%/91.33%。相比之下，direct exact-region projected-VSD scoring 并非在所有设置中最优，尤其在 LOMO 中不如 native VSD，因此 exact-region scoring 没有被设置为唯一 endpoint。

随后对完整正式三级路线进行端到端验证：projected-VSD Network Top3 beam generation，logCPM-compatible resolution-group reranking，以及 logCPM-compatible exact-region reranking。该路线在 LOSO 中的 Network Top3 为 92.38%，在 LOMO 中为 91.21%。Resolution-group Top3 分别为 72.36% 和 69.09%，而 exact-region Top3 分别为 45.33% 和 42.36%。这种随解剖分辨率升高而下降的准确率梯度支持把 Network Top3 作为主 endpoint，把 resolution group 作为更稳健的 region-level 输出；exact-region 结果保留为探索性局部候选排序，而不是确定性定位。

外部分析按照标签能支持的分辨率解释。Allen Human Brain Atlas 映射标签验证中，hybrid 路线的 Network Top1/Top3 为 74.68%/94.42%，resolution-group Top1/Top3 为 36.26%/67.03%，exact Top1/Top3 为 24.18%/42.86%；hybrid exact Top3 高于 logCPM baseline 和 projected-VSD-only，支持“projected VSD 用于 broad beam、logCPM 用于 local reranking”的路线选择。TCGA/BraTS glioma tissue RNA-seq + MRI label 只支持 coarse anatomical consistency，因为 MRI 真值是人脑 atlas 标签，而不是 Bo2023 猕猴 exact-region ID。在该数据集中，hybrid Network Top3 为 40.00%，broad-anatomy Top3 为 64.62%。GSE189919 用于验证 projection feasibility，而不是 accuracy，因为缺少可计算 brain-origin accuracy 的解剖真值。

### 使用边界

软件导出 ranked candidates、matched-gene coverage、entropy、margin、route identifier 和 resolution-specific warning。只有在存在独立解剖真值时才计算 accuracy。无标签 biofluid 预测作为 transfer stress test 或假设生成候选排序报告，而不是临床定位。最终 release package 计划包含非商业访问、测试数据、稳定 URL、版本化源码和归档 DOI。

## Funding

基金信息将在投稿系统中声明。

## Conflict of Interest

无。

## Data availability

Bo2023 猕猴脑转录组图谱、Allen Human Brain Atlas、TCGA/BraTS 和 GEO 数据集均可从原始仓库获取，并受其原始访问条件约束。可公开的 processed evaluation tables 和 figure source data 已按带版本的公开 release 与归档准备。

## References

Bakas,S. *et al.* (2017) Advancing The Cancer Genome Atlas glioma MRI collections with expert segmentation labels and radiomic features. *Sci. Data*, **4**, 170117.

Bo,T. *et al.* (2023) Brain-wide and cell-specific transcriptomic insights into MRI-derived cortical morphology in macaque monkeys. *Nat. Commun.*, **14**, 1283.

Hawrylycz,M.J. *et al.* (2012) An anatomically comprehensive atlas of the adult human brain transcriptome. *Nature*, **489**, 391-399.

Vorperian,S.K. *et al.* (2022) Cell types of origin of the cell-free transcriptome. *Nat. Biotechnol.*, **40**, 855-861.

## 图注

**图 1. cfRNA-BrainTrace 三级路线与验证证据。** query 只在 Network Top3 beam 阶段投影到 Bo2023-like VSD 空间；下游 resolution-group 与 exact-region reranking 在 logCPM 兼容的局部表达空间中完成。内部验证支持 LOSO 和 LOMO 设置下较高的 Network Top3 accuracy，同时 resolution-group accuracy 持续高于 exact-region accuracy。AHBA 支持 mapped-label 外部验证，而 TCGA/BraTS 只支持 coarse anatomical consistency。

**Alt text：** 多面板图概述 cfRNA-BrainTrace 工作流与验证结果。流程图显示 query 预处理、projected-VSD Network Top3 beam 生成和 logCPM 局部重排序。柱状图显示内部验证 Network Top3 接近 92%，resolution-group Top3 约 69%-72%，exact-region Top3 约 42%-45%。外部面板显示 AHBA Network Top3 为 94.42%，TCGA/BraTS broad-anatomy Top3 为 64.62%。
"""

CN_SUPP_MD = """# cfRNA-BrainTrace 中文补充材料

## 补充方法

### S1 正式验证路线

所有验证均按照软件正式路线组织。query profile 先表示为 logCPM/logTPM 兼容表达空间，再仅在 Network scoring 阶段投影到 Bo2023-like VSD 空间。Top3 Networks 构成 candidate beam。随后 resolution-group 和 exact-region candidates 都在该 beam 内使用 logCPM 兼容局部表达空间重排序。该设计把 broad candidate generation 与 fine regional interpretation 分开。

### S2 内部验证设计

内部验证使用 Bo2023 猕猴脑参考。Leave-one-sample-out validation 测试在留出单个样本时能否恢复其标签；leave-one-monkey-out validation 每次留出一只猴的全部样本，用于测试 donor-level generalization。指标分别在 Network、resolution group 和 exact region 三个层级报告。Top1、Top3 和 median true-rank 用于区分单一标签命中与候选列表召回。

### S3 外部验证设计

外部验证按数据集标签分辨率解释。AHBA human brain RNA-seq 用于 mapped-label validation，因为其解剖标签可映射到 Network、resolution group 和部分 exact-region 标签。TCGA/BraTS glioma tissue RNA-seq + MRI-derived labels 只用于 coarse anatomical consistency，因为其真值是人脑影像标签，不是 Bo2023 猕猴 exact-region ID。GSE189919 用于测试外部矩阵是否可投影到模型基因空间；由于缺少 patient-level anatomical truth，不用于 accuracy estimation。

### S4 图表组织

补充图 S1 给出主路线与验证总结图。补充表 S1-S5 分别给出内部验证设计、内部验证结果、外部验证设计、外部验证结果以及图表索引。这些材料共同说明正式路线如何验证、每个验证设置可支持何种结论。

## 补充结果

### S1 内部路线选择

第一组内部分析测试 Network-level candidate generation。projected VSD 在 LOSO 中 Network Top1/Top3 为 58.00%/91.58%，在 LOMO 中为 53.72%/91.33%，在 Network Top3 上均高于 logCPM baseline 与 native VSD。Direct exact-region scoring 的结果更低且更不稳定，尤其在 LOMO 中表现不足。因此正式路线使用 projected VSD 生成 broad Network beam，并使用 logCPM 兼容表达进行下游 local reranking。

### S2 正式内部三级验证

完整 LOSO validation 中，正式路线的 Network、resolution-group 和 exact-region Top3 分别为 92.38%、72.36% 和 45.33%。完整 LOMO validation 中，对应 Top3 分别为 91.21%、69.09% 和 42.36%。Median true-rank 从 Network 到 exact-region 层级逐步增大，符合更细解剖分辨率下不确定性升高的预期。因此 resolution group 是更适合的 region-level endpoint，exact-region output 保留为候选排序。

### S3 外部验证结果

AHBA 中，hybrid route 的 Network Top1/Top3 为 74.68%/94.42%，resolution-group Top1/Top3 为 36.26%/67.03%，exact-region Top1/Top3 为 24.18%/42.86%。Hybrid exact Top3 高于 logCPM baseline 的 30.77% 和 projected-VSD-only 的 29.67%。TCGA/BraTS 中 hybrid Network Top3 为 40.00%，broad-anatomy Top3 为 64.62%，只支持 coarse anatomical consistency。GSE189919 覆盖 15,622/21,668 projector genes，即 72.10% gene-space coverage，支持 projection feasibility，而不是 source-localization accuracy。
"""

CN_CHECKLIST_MD = """# Bioinformatics Application Note 中文投稿检查清单

官网要求核对日期：2026-06-19。

## 本版已落实

- 按 Application Note 写法组织为短摘要与正文。
- 摘要包含四个规定小标题：Summary、Availability and Implementation、Contact、Supplementary information。
- 主文不放主表，仅保留一张 Figure 1。
- Figure 1 已嵌入 DOCX，并单独导出低分辨率 PNG 与高分辨率 TIFF。
- Figure 1 按双栏宽 178 mm / 7.0 in 设计，白底。
- 图注后加入 Alt text。
- 补充方法、补充结果和补充表格整合为单个 PDF。
- 最新验证路线已改为 projected-VSD Network Top3 beam + logCPM downstream rerank。

## 待最终投稿包确认

- 作者、单位、通讯邮箱和基金信息。
- 公开 GitHub 或其他对非商业用户免费开放的稳定 URL。
- tagged release、测试数据和 Zenodo/Figshare 等归档 DOI。
- 在线应用 URL。
- OSI 认可许可证和维护承诺。
- 根据期刊政策决定是否披露 AI/工具辅助写作或代码生成。
"""

TABLES = {
    "TableS1_internal_validation_design.csv": [
        ["Validation", "Data", "Held-out unit", "Route tested", "Reported endpoints"],
        ["Network LOSO", "Bo2023 macaque brain RNA-seq", "Single sample", "Projected VSD vs logCPM/native VSD Network scoring", "Network Top1, Network Top3, median true-rank"],
        ["Network LOMO", "Bo2023 macaque brain RNA-seq", "One monkey", "Projected VSD vs logCPM/native VSD Network scoring", "Network Top1, Network Top3, median true-rank"],
        ["Formal three-tier LOSO", "Bo2023 macaque brain RNA-seq", "Single sample", "Projected-VSD Network beam plus logCPM group/exact rerank", "Network, resolution-group and exact-region Top1/Top3"],
        ["Formal three-tier LOMO", "Bo2023 macaque brain RNA-seq", "One monkey", "Fold-local Network, group and exact components", "Network, resolution-group and exact-region Top1/Top3"],
    ],
    "TableS2_internal_validation_results.csv": [
        ["Dataset", "Route", "Endpoint", "Top1", "Top3", "Median true-rank", "Interpretation"],
        ["Bo2023 LOSO", "Projected VSD Network", "Network", "58.00%", "91.58%", "1.0", "Supports projected-VSD Network beam"],
        ["Bo2023 LOMO", "Projected VSD Network", "Network", "53.72%", "91.33%", "1.0", "Supports donor-level Network beam"],
        ["Bo2023 LOSO", "Formal hybrid", "Network", "58.35%", "92.38%", "1.0", "Primary endpoint"],
        ["Bo2023 LOSO", "Formal hybrid", "Resolution group", "44.47%", "72.36%", "2.0", "Main region-level endpoint"],
        ["Bo2023 LOSO", "Formal hybrid", "Exact region", "22.48%", "45.33%", "4.0", "Exploratory ranking"],
        ["Bo2023 LOMO", "Formal hybrid", "Network", "57.75%", "91.21%", "1.0", "Cross-monkey support"],
        ["Bo2023 LOMO", "Formal hybrid", "Resolution group", "41.38%", "69.09%", "2.0", "Main region-level endpoint"],
        ["Bo2023 LOMO", "Formal hybrid", "Exact region", "22.17%", "42.36%", "5.0", "Exploratory ranking"],
    ],
    "TableS3_external_validation_design.csv": [
        ["Dataset", "Sample type", "n", "Truth label type", "Allowed conclusion"],
        ["AHBA", "Human normal brain RNA-seq", "242 total; 233 supported; 91 exact-evaluable", "Mapped anatomical labels", "Cross-species mapped-label validation"],
        ["TCGA/BraTS", "Glioma tissue RNA-seq with MRI-derived labels", "65 patients", "Human imaging labels", "Coarse anatomical consistency only"],
        ["GSE189919", "External count matrix", "51 samples", "No patient-level anatomical truth", "Projection feasibility only"],
    ],
    "TableS4_external_validation_results.csv": [
        ["Dataset", "Route", "Endpoint", "Top1", "Top3", "Conclusion"],
        ["AHBA", "Hybrid", "Network", "74.68%", "94.42%", "Strong mapped-label Network support"],
        ["AHBA", "Hybrid", "Resolution group", "36.26%", "67.03%", "Moderate group-level support"],
        ["AHBA", "Hybrid", "Exact region", "24.18%", "42.86%", "Only exact-evaluable mapped labels"],
        ["AHBA", "logCPM baseline", "Exact region", "17.58%", "30.77%", "Below hybrid exact Top3"],
        ["AHBA", "Projected VSD only", "Exact region", "10.99%", "29.67%", "Below hybrid exact Top3"],
        ["TCGA/BraTS", "Hybrid", "Network", "15.38%", "40.00%", "Coarse consistency only"],
        ["TCGA/BraTS", "Hybrid", "Broad anatomy", "13.85%", "64.62%", "Coarse consistency only"],
        ["GSE189919", "Projection feasibility", "Projector gene overlap", "15622/21668", "72.10%", "No accuracy claim"],
    ],
    "TableS5_figure_table_index.csv": [
        ["Item", "Content", "Use in manuscript"],
        ["Figure 1", "Formal route diagram and validation summary", "Main Application Note figure"],
        ["Supplementary Figure S1", "Main figure source artwork reproduced in the supplementary PDF", "Supplementary visual reference"],
        ["Table S1", "Internal validation design", "Documents LOSO/LOMO operations"],
        ["Table S2", "Internal validation results", "Reports Network, group and exact-region metrics"],
        ["Table S3", "External validation design", "Documents label support and allowed conclusions"],
        ["Table S4", "External validation results", "Reports AHBA, TCGA/BraTS and GSE189919 outcomes"],
    ],
    "TableS6_claim_boundaries.csv": [
        ["Avoid", "Use"],
        ["Projected VSD creates a new Bo2023 atlas.", "Only query profiles are projected for Network beam generation."],
        ["Projected VSD is best for exact-region inference.", "Projected VSD supports Network beam; logCPM supports downstream reranking."],
        ["TCGA/BraTS validates Bo2023 exact regions.", "TCGA/BraTS supports coarse anatomical consistency."],
        ["GSE189919 validates accuracy.", "GSE189919 verifies projection feasibility only."],
        ["Exact Top1 is deterministic localization.", "Exact outputs are exploratory local candidate rankings."],
    ],
}


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    for old_table in TABLE_DIR.glob("*.csv"):
        old_table.unlink()


def write_sources() -> None:
    (OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_20260619.md").write_text(MAIN_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_Supplement_submission_20260619.md").write_text(SUPP_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_Submission_Checklist_20260619.md").write_text(CHECKLIST_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_CN_20260619.md").write_text(CN_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_Supplement_submission_CN_20260619.md").write_text(CN_SUPP_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_Submission_Checklist_CN_20260619.md").write_text(CN_CHECKLIST_MD, encoding="utf-8")
    for name, rows in TABLES.items():
        with (TABLE_DIR / name).open("w", encoding="utf-8-sig", newline="") as handle:
            csv.writer(handle).writerows(rows)


def make_main_figure() -> None:
    width_px, height_px = 8400, 6000
    img = PILImage.new("RGB", (width_px, height_px), "white")
    draw = ImageDraw.Draw(img)

    def font(size: int, bold: bool = False):
        candidates = [
            r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return ImageFont.truetype(candidate, size)
        return ImageFont.load_default()

    title_font = font(170, True)
    panel_font = font(125, True)
    label_font = font(92)
    small_font = font(72)
    axis_font = font(82)
    value_font = font(74)

    def centered_text(xy, text, fnt, fill="#111827", spacing=18):
        lines = text.split("\n")
        heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
        total_h = sum(heights) + spacing * (len(lines) - 1)
        y = xy[1] - total_h / 2
        for line, h in zip(lines, heights):
            bbox = draw.textbbox((0, 0), line, font=fnt)
            draw.text((xy[0] - (bbox[2] - bbox[0]) / 2, y), line, font=fnt, fill=fill)
            y += h + spacing

    draw.text((width_px // 2 - draw.textbbox((0, 0), "cfRNA-BrainTrace three-tier Application Note summary", font=title_font)[2] // 2, 120),
              "cfRNA-BrainTrace three-tier Application Note summary", font=title_font, fill="#111827")

    draw.text((360, 560), "A", font=panel_font, fill="#111827")
    boxes = [
        ("Query\nlogCPM/logTPM", 720, "#EAF2F8"),
        ("Project to\nBo2023-like VSD", 2300, "#E8F5E9"),
        ("Network\nTop3 beam", 3880, "#E3F2FD"),
        ("logCPM local\nrerank", 5440, "#FFF3E0"),
        ("Network + group\n+ exact candidates", 7040, "#F3E5F5"),
    ]
    box_w, box_h = 1160, 520
    for text, cx, fill in boxes:
        x0, y0 = cx - box_w // 2, 820
        x1, y1 = cx + box_w // 2, 820 + box_h
        draw.rounded_rectangle([x0, y0, x1, y1], radius=55, fill=fill, outline="#4B5563", width=8)
        centered_text((cx, (y0 + y1) / 2), text, label_font)
    for (_, cx1, _), (_, cx2, _) in zip(boxes[:-1], boxes[1:]):
        y = 1080
        draw.line((cx1 + box_w // 2 + 35, y, cx2 - box_w // 2 - 75, y), fill="#111827", width=10)
        draw.polygon([(cx2 - box_w // 2 - 75, y - 35), (cx2 - box_w // 2 - 75, y + 35), (cx2 - box_w // 2 - 20, y)], fill="#111827")
    note = "Projected VSD is restricted to Network beam generation; downstream region-level reranking uses logCPM-compatible local expression."
    draw.text((520, 1480), note, font=small_font, fill="#374151")

    draw.text((360, 1880), "B", font=panel_font, fill="#111827")
    plot_x0, plot_y0, plot_x1, plot_y1 = 780, 2250, 8020, 5200
    draw.line((plot_x0, plot_y1, plot_x1, plot_y1), fill="#111827", width=8)
    draw.line((plot_x0, plot_y0, plot_x0, plot_y1), fill="#111827", width=8)
    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = plot_y1 - int((plot_y1 - plot_y0) * tick)
        draw.line((plot_x0 - 30, y, plot_x0, y), fill="#111827", width=6)
        draw.line((plot_x0, y, plot_x1, y), fill="#E5E7EB", width=3)
        draw.text((plot_x0 - 170, y - 45), f"{int(tick * 100)}%", font=axis_font, fill="#111827")
    draw.text((plot_x0, 2130), "Top3 accuracy / coverage (%)", font=axis_font, fill="#111827")

    categories = ["Internal\nLOSO", "Internal\nLOMO", "AHBA\nexternal", "TCGA/BraTS\nMRI"]
    series = [
        ("Network Top3", [0.9238, 0.9121, 0.9442, 0.4000], "#1f77b4"),
        ("Resolution group Top3", [0.7236, 0.6909, 0.6703, None], "#2ca02c"),
        ("Exact region Top3", [0.4533, 0.4236, 0.4286, None], "#ff9f1c"),
        ("Broad anatomy Top3", [None, None, None, 0.6462], "#7b3294"),
    ]
    group_centers = [1500, 3200, 4900, 6600]
    bar_w = 160
    offsets = [-270, -90, 90, 270]
    for s_idx, (name, vals, color) in enumerate(series):
        lx = 2250 + s_idx * 1550
        draw.rectangle((lx, 1960, lx + 140, 2035), fill=color)
        draw.text((lx + 170, 1950), name, font=small_font, fill="#111827")
        for cx, value in zip(group_centers, vals):
            if value is None:
                continue
            x0 = cx + offsets[s_idx] - bar_w // 2
            y0 = plot_y1 - int((plot_y1 - plot_y0) * value)
            draw.rectangle((x0, y0, x0 + bar_w, plot_y1), fill=color)
            text = f"{value * 100:.1f}%"
            tw = draw.textbbox((0, 0), text, font=value_font)[2]
            draw.text((x0 + bar_w / 2 - tw / 2, y0 - 90), text, font=value_font, fill="#111827")
    for cx, label in zip(group_centers, categories):
        centered_text((cx, 5480), label, axis_font)
    draw.text((4020, 5780), "TCGA/BraTS reports coarse anatomical consistency only; no Bo2023 exact-region accuracy is claimed.", font=small_font, fill="#374151")

    low_png = FIG_DIR / "Figure1_cfRNA_BrainTrace_Bioinformatics_lowres.png"
    high_tif = FIG_DIR / "Figure1_cfRNA_BrainTrace_Bioinformatics_highres_178mm.tif"
    img.save(high_tif, dpi=(1200, 1200), compression="tiff_lzw")
    img.resize((2100, 1500), PILImage.Resampling.LANCZOS).save(low_png, dpi=(300, 300))


def set_font(run, size=10, bold=None, italic=None, name="Arial", east_asia="Arial", color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown(paragraph, text, size=10, east_asia="Arial"):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size, east_asia=east_asia)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, east_asia=east_asia)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True, east_asia=east_asia)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, name="Consolas", east_asia=east_asia)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size, east_asia=east_asia)


def build_docx(source_md: str, output_name: str, cn: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if cn else "Arial")
    normal.font.size = Pt(9.5 if not cn else 10.5)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(4)

    for name, size, before, after in [("Heading 1", 13, 8, 3), ("Heading 2", 11, 6, 2), ("Heading 3", 10, 4, 1)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei" if cn else "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    lines = source_md.splitlines()
    in_legend = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line in {"## Figure legend", "## 图注"}:
            in_legend = True
            doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run("Figure 1" if not cn else "图 1")
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = pic.add_run().add_picture(str(FIG_DIR / "Figure1_cfRNA_BrainTrace_Bioinformatics_lowres.png"), width=Inches(7.0))
            shape._inline.docPr.set("title", "Figure 1. cfRNA-BrainTrace three-tier route and validation evidence")
            shape._inline.docPr.set("descr", "Workflow and bar chart summarizing three-tier validation results.")
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown(p, line[2:], size=15, east_asia="SimHei" if cn else "Arial")
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(23, 54, 93)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_markdown(p, line[4:], size=10.5, east_asia="SimHei" if cn else "Arial")
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_markdown(p, line[3:], size=12.5, east_asia="SimHei" if cn else "Arial")
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not in_legend else WD_ALIGN_PARAGRAPH.LEFT
            add_markdown(p, line, size=10.5 if cn else 9.5, east_asia="SimSun" if cn else "Arial")

    core = doc.core_properties
    core.title = "cfRNA-BrainTrace Bioinformatics Application Note submission draft" if not cn else "cfRNA-BrainTrace Bioinformatics Application Note 中文投稿稿"
    core.subject = "Application Note"
    core.comments = "Built against Bioinformatics Application Note constraints on 2026-06-19."
    out_path = OUT / output_name
    tmp_path = OUT / f".{output_name}.tmp.docx"
    doc.save(tmp_path)
    try:
        tmp_path.replace(out_path)
    except PermissionError:
        fallback = OUT / output_name.replace(".docx", "_updated.docx")
        tmp_path.replace(fallback)
        print(f"WARNING: {out_path.name} is locked; wrote {fallback.name}")


def register_font() -> str:
    for candidate in [r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\simhei.ttf"]:
        if Path(candidate).exists():
            pdfmetrics.registerFont(TTFont("DocFont", candidate))
            return "DocFont"
    return "Helvetica"


def fit_image(path: Path, max_w: float, max_h: float) -> Image:
    with PILImage.open(path) as img:
        width, height = img.size
    scale = min(max_w / width, max_h / height)
    return Image(str(path), width=width * scale, height=height * scale)


def pdf_cell(text: str, style) -> Paragraph:
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def build_pdf_table(path: Path, content_w: float, body, header) -> Table:
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    data = [[pdf_cell(v, header) for v in rows[0]]]
    data.extend([[pdf_cell(v, body) for v in row] for row in rows[1:]])
    widths = [content_w / len(rows[0])] * len(rows[0])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FA")]),
    ]))
    return table


def build_supplement_pdf(source_md: str = SUPP_MD, output_name: str = "Bioinformatics_Application_Note_Supplementary_File_submission_20260619.pdf", cn: bool = False) -> None:
    font = register_font()
    page = landscape(A4)
    page_w, page_h = page
    margin = 12 * mm
    content_w = page_w - 2 * margin
    content_h = page_h - 2 * margin
    styles = getSampleStyleSheet()
    title = ParagraphStyle("Title", parent=styles["Title"], fontName=font, fontSize=18, leading=23, alignment=TA_CENTER, textColor=colors.HexColor("#17365D"))
    subtitle = ParagraphStyle("Subtitle", parent=styles["BodyText"], fontName=font, fontSize=10, leading=14, alignment=TA_CENTER)
    heading = ParagraphStyle("Heading", parent=styles["Heading1"], fontName=font, fontSize=12, leading=15, textColor=colors.HexColor("#1F4E79"), spaceAfter=5)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=font, fontSize=8.5, leading=11)
    cell = ParagraphStyle("Cell", parent=styles["BodyText"], fontName=font, fontSize=7.2, leading=9)
    header = ParagraphStyle("Header", parent=cell, fontName=font, textColor=colors.white, alignment=TA_CENTER)
    doc = SimpleDocTemplate(
        str(OUT / output_name),
        pagesize=page,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )
    story = [
        Spacer(1, 25 * mm),
        Paragraph("cfRNA-BrainTrace Supplementary File" if not cn else "cfRNA-BrainTrace 中文补充材料", title),
        Paragraph("Single supplementary file with methods, results, figures and tables" if not cn else "方法、结果、图和表的单一补充文件", subtitle),
        Spacer(1, 8 * mm),
        Paragraph("Updated for Bioinformatics Application Note submission draft, 2026-06-19" if not cn else "Bioinformatics Application Note 中文投稿参考版，2026-06-19", subtitle),
        PageBreak(),
        Paragraph("Supplementary Methods and Results" if not cn else "补充方法与结果", heading),
    ]
    for line in source_md.splitlines():
        if line.startswith("### "):
            story.append(Paragraph(line[4:], heading))
        elif line and not line.startswith("#") and not line.startswith("##"):
            story.append(Paragraph(line, body))
    story.append(PageBreak())
    story.append(Paragraph("Supplementary Figure S1. Main figure source artwork" if not cn else "补充图 S1. 主图源图", heading))
    story.append(fit_image(FIG_DIR / "Figure1_cfRNA_BrainTrace_Bioinformatics_lowres.png", content_w, content_h - 25 * mm))
    story.append(PageBreak())
    for idx, table_path in enumerate(sorted(TABLE_DIR.glob("*.csv")), start=1):
        story.append(Paragraph(f"Supplementary Table {idx}. {table_path.stem}", heading))
        story.append(build_pdf_table(table_path, content_w, cell, header))
        if idx < len(list(TABLE_DIR.glob("*.csv"))):
            story.append(PageBreak())
    doc.build(story)


def word_count() -> int:
    text = re.sub(r"`[^`]+`", " code ", MAIN_MD)
    text = re.sub(r"[#*|]", " ", text)
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def main() -> None:
    ensure_dirs()
    write_sources()
    make_main_figure()
    build_docx(MAIN_MD, "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_20260619.docx")
    build_docx(CN_MD, "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_CN_20260619.docx", cn=True)
    build_supplement_pdf()
    build_supplement_pdf(CN_SUPP_MD, "Bioinformatics_Application_Note_Supplementary_File_submission_CN_20260619.pdf", cn=True)
    (OUT / "word_count.txt").write_text(f"Approximate manuscript word count: {word_count()}\n", encoding="utf-8")
    print(OUT)
    print(f"Approximate word count: {word_count()}")


if __name__ == "__main__":
    main()
