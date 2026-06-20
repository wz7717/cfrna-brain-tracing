from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "bioinformatics_application_note_submission_20260619"
FIGURE = OUT / "figures" / "Figure1_cfRNA_BrainTrace_Bioinformatics_lowres.png"
MD_OUT = OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_revised.md"
DOCX_OUT = OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_revised.docx"
CHANGELOG_OUT = OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_revised_changelog.md"


REVISED_MD = """# cfRNA-BrainTrace: hierarchical brain-origin candidate ranking from RNA expression profiles with a primate transcriptomic atlas

**Article type:** Application Note  
**Category:** Gene expression  
**Authors and affiliations:** Author information to be supplied in the submission system  
**Corresponding author:** [CORRESPONDING AUTHOR EMAIL REQUIRED BEFORE SUBMISSION]

## Abstract

### Summary

cfRNA-BrainTrace is a Python and Streamlit application for hierarchical brain-origin candidate ranking from RNA expression profiles. The locked route projects query profiles into a Bo2023-like variance-stabilized space to generate a broad 10-class macaque functional-anatomical Network Top3 beam, then uses logCPM-compatible local expression for resolution-group and exploratory exact-region reranking. The software reports ranked candidates, marker coverage, entropy, margins, route identifiers and scope warnings. In internal validation, Network Top3 accuracy reached 92.38% in leave-one-sample-out and 91.21% in leave-one-monkey-out evaluation, with lower accuracy at resolution-group and exact-region levels. AHBA mapped-label validation supported coarse cross-species transfer, whereas TCGA/BraTS and unlabeled biofluid analyses defined transfer limitations. cfRNA-BrainTrace is intended for reproducible coarse candidate ranking and resolution-limit auditing, not stand-alone clinical localization from unlabeled biofluid RNA.

### Availability and Implementation

Implemented in Python 3.11+ with command-line and Streamlit interfaces. The development repository is **https://github.com/wz7717/cfrna-brain-tracing** and contains README documentation, installation instructions, a command-line entry point, Streamlit entry points, tests, lightweight model artifacts and validation/benchmark scripts. The manuscript version is archived at **[ZENODO DOI REQUIRED BEFORE SUBMISSION]**. A live demonstration is available at **[STABLE WEB SERVICE URL REQUIRED BEFORE SUBMISSION]**. The software is released under **[LICENSE REQUIRED BEFORE SUBMISSION]**. Public example input/output files are **[EXAMPLE INPUT/OUTPUT REQUIRED BEFORE SUBMISSION]**.

### Contact

**[CORRESPONDING AUTHOR EMAIL REQUIRED BEFORE SUBMISSION]**

### Supplementary information

Supplementary data are available online.

## Text

### Motivation

RNA expression profiles can retain tissue-of-origin information, but within-brain source tracing is limited by transcriptional similarity among neighbouring regions, atlas granularity and domain shifts between reference tissue, tumour tissue and biofluid RNA. These constraints are central to atlas-based brain RNA tracing: a stable broad candidate may be recoverable even when exact-region localization is not. A single exact-region output can therefore imply more anatomical resolution than the data support. cfRNA-BrainTrace addresses this problem by returning a hierarchy of Network, resolution-group and exact-region candidates together with confidence, coverage and scope diagnostics.

### System and methods

The locked production route was selected from validation evidence rather than from a single modelling assumption. First, projected variance-stabilized expression representation was consistently useful for broad Network candidate generation, where Network denotes the 10-class macaque functional-anatomical source space used by the Bo2023 reference. Second, direct exact-region scoring in projected-VSD space was not consistently superior to native or logCPM-based alternatives. Third, local reranking within the Network beam improved the interpretability of downstream anatomical outputs. The resulting default scoring route is therefore a hybrid route: projected-VSD expression is used for broad candidate generation, while logCPM-compatible expression is used for downstream local ranking.

Operationally, cfRNA-BrainTrace normalizes the uploaded expression table into logCPM- or logTPM-compatible space and aligns genes to the reference panel. A reference-fitted linear projector maps the query to Bo2023-like VSD space using stored gene-wise slope, intercept and clipping parameters; the default route does not use target-cohort labels or target-cohort distribution information. The projected representation is used only for Network scoring. The top three Networks form a candidate beam, and candidate regions outside this beam are excluded from downstream regional ranking. Within the retained beam, resolution groups are reranked using local logCPM-compatible expression, and exact regions are then reranked as lower-confidence local candidates. This route uses the Bo2023 macaque brain transcriptomic reference without converting the atlas into a new projected atlas.

### Implementation

The Network projection and scoring route is implemented in `core/network_tracing.py`, and the three-tier regional route is implemented in `core/bo2023_region_tracing.py`. The command-line interface (`cli.py`) and Streamlit entry points (`streamlit_app.py`, `app/`) call the same scoring core to avoid interface-specific divergence. Model artifacts include marker order, Network centroids, anatomical dictionaries, route parameters and warning metadata. The repository includes unit tests under `tests/` and validation/benchmark scripts under `scripts/` and `benchmark_runner.py`, supporting reproducible reruns of the analyses reported here.

### Validation

Validation was matched to the resolution supported by available labels. Internal experiments first tested whether projected-VSD query representation was appropriate for generating the broad Network candidate beam. In fold-local leave-one-sample-out Network validation, projected VSD achieved Top1/Top3 accuracy of 58.00%/91.58%, exceeding logCPM baseline and native VSD routes. In strict leave-one-monkey-out validation, projected VSD reached 53.72%/91.33%. Direct exact-region projected-VSD scoring was lower and less stable, especially in leave-one-monkey-out evaluation, so exact-region scoring was not used as the sole endpoint.

The complete validation then tested the locked three-tier procedure end to end: projected-VSD Network Top3 beam generation, logCPM-compatible resolution-group reranking and logCPM-compatible exact-region reranking. This route achieved Network Top3 accuracy of 92.38% in leave-one-sample-out validation and 91.21% in leave-one-monkey-out validation. Resolution-group Top3 accuracy was 72.36% and 69.09%, respectively, whereas exact-region Top3 was 45.33% and 42.36%. The accuracy gradient across levels supports Network Top3 as the primary endpoint and resolution group as the more defensible region-level output. Exact-region results are retained as exploratory local candidate rankings rather than deterministic localization.

In AHBA mapped-label external validation, the hybrid route achieved Network Top1/Top3 accuracy of 74.68%/94.42%, resolution-group Top1/Top3 accuracy of 36.26%/67.03% and exact Top1/Top3 accuracy of 24.18%/42.86%; hybrid exact Top3 exceeded both logCPM baseline and projected-VSD-only scoring, supporting the use of projected VSD for the broad beam and logCPM expression for local reranking. In TCGA/BraTS glioma tissue RNA-seq with MRI-derived labels, results support only coarse anatomical consistency because MRI truth labels are human atlas labels rather than Bo2023 macaque exact-region identifiers. In that setting, hybrid Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%. GSE189919 and other biofluid datasets without anatomical truth were treated as projection-feasibility or transfer stress tests rather than localization-accuracy validation.

### Use and limitations

The intended use of cfRNA-BrainTrace is coarse brain-origin candidate ranking from RNA expression profiles and assessment of whether a sample supports Network-level, resolution-group-level or only low-confidence output. The software exports warnings for sparse input profiles, low marker coverage, high entropy, low score margins, out-of-scope anatomical spaces and domain shift. It is not intended for deterministic exact-region localization, clinical cfRNA localization without independent anatomical truth, or localization of cerebellar/posterior-fossa samples outside the current reference space. The cfRNA name reflects the intended prospective research setting; the current evidence does not establish clinical liquid-biopsy localization.

## Funding

[FUNDING INFORMATION REQUIRED BEFORE SUBMISSION]

## Conflict of Interest

None declared.

## Data availability

The Bo2023 macaque transcriptomic atlas, Allen Human Brain Atlas, TCGA/BraTS and GEO datasets are available from their original repositories under their original access conditions. Processed non-sensitive evaluation tables and figure source data will be archived at **[DOI REQUIRED BEFORE SUBMISSION]** and must be available before submission.

## References

Bakas,S. *et al.* (2017) Advancing The Cancer Genome Atlas glioma MRI collections with expert segmentation labels and radiomic features. *Sci. Data*, **4**, 170117.

Bo,T. *et al.* (2023) Brain-wide and cell-specific transcriptomic insights into MRI-derived cortical morphology in macaque monkeys. *Nat. Commun.*, **14**, 1283.

Hawrylycz,M.J. *et al.* (2012) An anatomically comprehensive atlas of the adult human brain transcriptome. *Nature*, **489**, 391-399.

Vorperian,S.K. *et al.* (2022) Cell types of origin of the cell-free transcriptome. *Nat. Biotechnol.*, **40**, 855-861.

## Figure legend

**Figure 1. cfRNA-BrainTrace three-tier route and validation evidence.** The query is projected into Bo2023-like VSD space only for Network Top3 beam generation; downstream resolution-group and exact-region reranking is performed in logCPM-compatible local expression space. Internal validation shows the expected resolution gradient, with high Network Top3 accuracy and lower resolution-group and exact-region accuracy. AHBA provides mapped-label external validation, whereas TCGA/BraTS supports only coarse anatomical consistency because its MRI truth labels are human atlas labels rather than Bo2023 macaque exact-region identifiers.

**Alt text:** Multi-panel figure summarizing the cfRNA-BrainTrace workflow and validation. A workflow diagram shows query preprocessing, projected-VSD Network Top3 beam generation and logCPM local reranking. Bar charts show Network Top3 near 92% in internal validation, resolution-group Top3 around 69%-72% and lower exact-region Top3 around 42%-45%. External panels show AHBA Network Top3 of 94.42% and TCGA/BraTS broad-anatomy Top3 of 64.62%.
"""


CHANGELOG = """# Revised Application Note change log

## Main wording changes

- Retitled the manuscript to emphasize brain-origin candidate ranking from RNA expression profiles rather than validated clinical cfRNA localization.
- Rewrote the Summary to state the locked route, main validation results and clinical/localization boundary.
- Replaced future-tense availability claims with verified repository contents plus explicit placeholders for DOI, web service URL, license, example input/output and corresponding author email.
- Defined Network as the 10-class macaque functional-anatomical source space and projected-VSD as projected variance-stabilized expression representation.
- Clarified that the reference-fitted linear projector uses stored slope/intercept/clipping parameters and does not use target-cohort labels or target-cohort distribution information in the default route.
- Added a separate Implementation section covering `core/network_tracing.py`, `core/bo2023_region_tracing.py`, `cli.py`, Streamlit entry points, model artifacts, tests and validation/benchmark scripts.
- Reframed validation around label-supported resolution: internal Network beam, formal three-tier route, AHBA mapped-label validation, TCGA/BraTS coarse consistency and GSE189919/biofluid transfer stress tests.
- Strengthened Use and limitations, including sparse profiles, marker coverage, entropy, margin, out-of-scope anatomy and domain shift warnings.

## Remaining placeholders

- `[CORRESPONDING AUTHOR EMAIL REQUIRED BEFORE SUBMISSION]`
- `[ZENODO DOI REQUIRED BEFORE SUBMISSION]`
- `[STABLE WEB SERVICE URL REQUIRED BEFORE SUBMISSION]`
- `[LICENSE REQUIRED BEFORE SUBMISSION]`
- `[EXAMPLE INPUT/OUTPUT REQUIRED BEFORE SUBMISSION]`
- `[FUNDING INFORMATION REQUIRED BEFORE SUBMISSION]`
- `[DOI REQUIRED BEFORE SUBMISSION]`

## Repository items found

- README documentation and installation/run instructions: `README.md`
- Streamlit entry points: `streamlit_app.py`, `cfrna_tracing_app.py`, `app/`
- Command-line interface: `cli.py`; package entry point `cfrna-tracing=cli:main` in `setup.py`
- Tests: `tests/`
- Lightweight model artifacts: `data/models/`
- Validation and benchmark scripts: `scripts/`, `benchmark_runner.py`

## Repository items not found or not verified

- Root-level OSI-approved license file
- Git tag or tagged release
- Zenodo/Figshare archival DOI
- Confirmed live Streamlit/web service URL
- Public example input/output files suitable for submission

## Validation numbers

No validation numbers were changed. The revised draft retains the provided values for internal projected-VSD Network validation, formal three-tier LOSO/LOMO validation, AHBA mapped-label validation and TCGA/BraTS coarse anatomical consistency.
"""


def set_font(run, size=10, bold=None, italic=None, name="Arial", color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown(paragraph, text, size=10):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, name="Consolas")
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size)


def build_docx(markdown: str, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(4)

    for name, size, before, after in [("Heading 1", 13, 8, 3), ("Heading 2", 11, 6, 2), ("Heading 3", 10, 4, 1)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    in_legend = False
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line == "## Figure legend":
            in_legend = True
            doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run("Figure 1")
            if FIGURE.exists():
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shape = pic.add_run().add_picture(str(FIGURE), width=Inches(7.0))
                shape._inline.docPr.set("title", "Figure 1. cfRNA-BrainTrace three-tier route and validation evidence")
                shape._inline.docPr.set("descr", "Workflow and bar chart summarizing three-tier validation results.")
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown(p, line[2:], size=15)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(23, 54, 93)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_markdown(p, line[4:], size=10.5)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_markdown(p, line[3:], size=12.5)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if in_legend else WD_ALIGN_PARAGRAPH.JUSTIFY
            add_markdown(p, line, size=9.5)

    core = doc.core_properties
    core.title = "cfRNA-BrainTrace Bioinformatics Application Note revised draft"
    core.subject = "Bioinformatics Application Note"
    core.comments = "Revised according to repository-supported submission guidance."
    doc.save(output)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text(REVISED_MD, encoding="utf-8")
    CHANGELOG_OUT.write_text(CHANGELOG, encoding="utf-8")
    build_docx(REVISED_MD, DOCX_OUT)
    words = len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", re.sub(r"`[^`]+`", " code ", REVISED_MD)))
    print(DOCX_OUT)
    print(MD_OUT)
    print(CHANGELOG_OUT)
    print(f"Approximate word count: {words}")


if __name__ == "__main__":
    main()
