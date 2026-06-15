from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "Bioinformatics_Application_Note_cfRNA_BrainTrace_20260614.md"
FIGURE = ROOT / "manuscript" / "figures_application_note_20260614" / "Figure1_cfRNA_BrainTrace_application_note.png"
TABLE = ROOT / "manuscript" / "tables_application_note_20260614" / "Table1_cfRNA_BrainTrace_features.csv"
OUTPUT = ROOT / "manuscript" / "Bioinformatics_Application_Note_cfRNA_BrainTrace_20260614.docx"


def set_font(run, size=10, bold=None, italic=None, name="Arial", color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown(paragraph, text, size=10):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, name="Consolas")
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10)
normal.paragraph_format.line_spacing = 1.08
normal.paragraph_format.space_after = Pt(3)

for name, size, before, after in [
    ("Heading 1", 12.5, 8, 3),
    ("Heading 2", 11, 6, 2),
    ("Heading 3", 10, 4, 1),
]:
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(31, 78, 121)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("BIOINFORMATICS | APPLICATION NOTE"), size=8, bold=True, color="6B7280")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

lines = SOURCE.read_text(encoding="utf-8").splitlines()
for raw in lines:
    line = raw.strip()
    if not line:
        continue
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        add_markdown(p, line[2:], size=16)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(23, 54, 93)
        continue
    if line.startswith("## Figure legend"):
        break
    if line.startswith("## "):
        p = doc.add_paragraph(style="Heading 1")
        add_markdown(p, line[3:], size=12.5)
        continue
    if line.startswith("### "):
        p = doc.add_paragraph(style="Heading 2")
        add_markdown(p, line[4:], size=11)
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    add_markdown(p, line, size=10)

doc.add_paragraph(style="Heading 1").add_run("Figure 1")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = p.add_run().add_picture(str(FIGURE), width=Inches(6.55))
shape._inline.docPr.set(
    "descr",
    "cfRNA-BrainTrace software workflow, internal validation, external coarse-resolution evaluation and interpretation policy.",
)
shape._inline.docPr.set("title", "Figure 1. cfRNA-BrainTrace workflow and validation")
caption = doc.add_paragraph()
caption.paragraph_format.space_after = Pt(5)
add_markdown(
    caption,
    "Fig. 1. cfRNA-BrainTrace workflow and validation. "
    "(A) Software inputs, locked scoring route and hierarchical outputs. "
    "(B) Network Top1 and Top3 accuracy in LOSO and LOMO validation. "
    "(C) External coarse-resolution performance in normal human brain and paired glioma data. "
    "(D) Accuracy is reported only with independent anatomical truth; unlabeled biofluids are transfer stress tests.",
    size=8.5,
)

import csv

rows = list(csv.reader(TABLE.open(encoding="utf-8-sig")))
doc.add_paragraph(style="Heading 1").add_run("Table 1")
caption = doc.add_paragraph()
add_markdown(caption, "Table 1. Principal software functions and outputs.", size=8.5)
table = doc.add_table(rows=len(rows), cols=len(rows[0]))
table.autofit = False
widths = [1100, 4100, 4100]
for i, row in enumerate(rows):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        cell.width = Inches(widths[j] / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        set_font(para.add_run(value), size=7.5, bold=(i == 0))
        if i == 0:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9EAF7")
            cell._tc.get_or_add_tcPr().append(shading)
if table.rows:
    repeat_header(table.rows[0])

core = doc.core_properties
core.title = "cfRNA-BrainTrace: hierarchical brain-origin inference from RNA-seq with a primate transcriptomic atlas"
core.subject = "Bioinformatics Application Note"
core.author = ""
core.comments = "Converted to Application Note format on 2026-06-14; submission placeholders remain."
core.keywords = "cfRNA; RNA-seq; brain origin; transcriptomic atlas; Streamlit; Application Note"

doc.save(OUTPUT)
print(OUTPUT)
