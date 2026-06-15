from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "Bioinformatics_Application_Note_cfRNA_BrainTrace_CN_20260614.md"
FIGURE = ROOT / "manuscript" / "figures_application_note_20260614" / "Figure1_cfRNA_BrainTrace_application_note.png"
OUTPUT = ROOT / "manuscript" / "Bioinformatics_Application_Note_cfRNA_BrainTrace_CN_20260614.docx"


def set_font(run, size=10.5, bold=None, italic=None, latin="Arial", east_asia="宋体", color=None):
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown(paragraph, text, size=10.5):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, latin="Consolas", east_asia="宋体")
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.35
normal.paragraph_format.space_after = Pt(4)

for name, size, before, after in [
    ("Heading 1", 14, 10, 4),
    ("Heading 2", 12, 8, 3),
    ("Heading 3", 11, 6, 2),
]:
    style = doc.styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(31, 78, 121)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("BIOINFORMATICS APPLICATION NOTE | 中文版"), size=8, bold=True, color="6B7280")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

lines = SOURCE.read_text(encoding="utf-8").splitlines()
in_table = False
table_rows = []

for raw in lines:
    line = raw.strip()
    if not line:
        continue
    if line.startswith("|"):
        in_table = True
        if "---" not in line:
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
        continue
    if in_table:
        in_table = False
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        add_markdown(p, line[2:], size=17)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(23, 54, 93)
        continue
    if line.startswith("## 图注") or line.startswith("## 表格"):
        continue
    if line.startswith("## "):
        p = doc.add_paragraph(style="Heading 1")
        add_markdown(p, line[3:], size=14)
        continue
    if line.startswith("### "):
        p = doc.add_paragraph(style="Heading 2")
        add_markdown(p, line[4:], size=12)
        continue
    if line.startswith("**图1."):
        doc.add_paragraph(style="Heading 1").add_run("图1")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(FIGURE), width=Inches(6.45))
        shape._inline.docPr.set("descr", "cfRNA-BrainTrace软件工作流程、内部验证、外部粗粒度评价及结果解释原则。")
        shape._inline.docPr.set("title", "图1. cfRNA-BrainTrace工作流程与验证")
        caption = doc.add_paragraph()
        add_markdown(caption, line, size=9)
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    add_markdown(p, line, size=10.5)

if table_rows:
    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table.autofit = False
    widths = [1300, 4000, 4000]
    for i, row in enumerate(table_rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.1
            set_font(para.add_run(value), size=8.5, bold=(i == 0))
            if i == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shading)
    repeat_header(table.rows[0])

core = doc.core_properties
core.title = "cfRNA-BrainTrace：基于灵长类转录组图谱的RNA-seq分层脑来源推断"
core.subject = "Bioinformatics Application Note 中文版"
core.author = ""
core.comments = "中文版生成于2026-06-14；保留投稿前信息占位符。"
core.keywords = "cfRNA; RNA-seq; 脑来源; 转录组图谱; Streamlit; Application Note"

doc.save(OUTPUT)
print(OUTPUT)
