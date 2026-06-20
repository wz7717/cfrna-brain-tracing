from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from PIL import Image as PILImage


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "application_note_revision_20260619"
FIG_OUT = OUT / "figures"
TAB_OUT = OUT / "tables"

RESULTS = ROOT / "results" / "bo2023_reference_projection_20260616_cleaned_symbols"


MAIN_MD = """# cfRNA-BrainTrace: hierarchical brain-origin inference from RNA-seq with a primate transcriptomic atlas

**Article type:** Application Note  
**Category:** Gene expression  
**Authors and affiliations:** [TO BE INSERTED BEFORE SUBMISSION]  
**Corresponding author:** [NAME AND EMAIL TO BE INSERTED BEFORE SUBMISSION]

## Abstract

### Summary

cfRNA-BrainTrace is a Python and Streamlit application for hierarchical brain-origin candidate ranking from RNA expression profiles. The current locked route maps each query from logCPM- or logTPM-compatible expression into Bo2023-like variance-stabilized space only for 10-class Network candidate generation, then reranks resolution-group and exact-region candidates in local logCPM-compatible expression space. This design avoids treating VSD projection as a new atlas and exposes anatomical resolution as an explicit output. In Bo2023 internal validation, the formal three-tier hybrid route achieved Network Top3 accuracy of 0.9238 in leave-one-sample-out validation and 0.9121 in leave-one-monkey-out validation. Resolution-group Top3 accuracy was 0.7236 and 0.6909, respectively, whereas exact-region Top3 remained lower at 0.4533 and 0.4236. In AHBA cross-species mapped-label validation, the hybrid route retained high Network Top3 accuracy (0.9442) and improved exact Top3 accuracy (0.4286) over both logCPM baseline and projected-VSD-only routes. TCGA/BraTS glioma tissue RNA-seq with MRI labels supported only coarse anatomical consistency. The software therefore provides reproducible hierarchical candidate ranking, confidence diagnostics and claim boundaries for atlas-based RNA source tracing.

### Availability and Implementation

Python 3.11+; command-line and Streamlit interfaces. Source code and documentation are hosted in a private development repository at **https://github.com/wz7717/cfrna-brain-tracing** and require public release, archival DOI and licence confirmation before submission. A stable live application URL is **[TO BE INSERTED BEFORE SUBMISSION]**.

### Contact

**[CORRESPONDING AUTHOR EMAIL REQUIRED BEFORE SUBMISSION]**

### Supplementary information

Supplementary data are available online.

## 1 Introduction

Cell-free and tissue RNA profiles can retain tissue-of-origin information, but anatomical inference within the brain is constrained by shared transcriptional programmes, atlas resolution and sample-domain shifts. A single exact-region label can therefore be misleading when the evidence supports only a broader candidate set. This is especially relevant for translational studies in which coarse anatomical consistency may be reproducible, while exact-region localization remains underdetermined.

cfRNA-BrainTrace implements a hierarchical route using the Bo2023 macaque brain transcriptomic atlas and a versioned anatomical dictionary. The software reports Network, resolution-group and exact-region candidates with coverage, entropy, margin and scope diagnostics. The principal endpoint is Network Top3 candidate generation; resolution groups are the main region-level reporting granularity; exact-region results are interpreted as local candidate rankings rather than deterministic calls.

## 2 Software description

### 2.1 Inputs and preprocessing

The application accepts tabular gene-expression input with gene symbols and abundance values. Metadata columns are preserved in exports when supplied. Gene identifiers are normalized, audited against the reference panel and intersected with model genes. External count-like or TPM-like matrices are converted to logCPM/logTPM-compatible space before downstream scoring. Matched-gene counts, non-zero coverage and out-of-scope warnings are reported for each sample.

### 2.2 Reference projection and hierarchical inference

The training reference remains the original Bo2023 expression resource: VSD/batch-removed expression for Network candidate generation and raw-feature-count-derived logCPM expression for local downstream reranking. A linear projector maps held-out or external query profiles into Bo2023-like VSD space for Network Top3 beam generation only. The projector fitted the Bo2023 training samples well at the sample level (global Pearson 0.9961; median sample Pearson 0.9969), but gene-level reconstruction was weaker, so the software does not claim precise VSD recovery for every gene.

After the projected-VSD Network Top3 beam is generated, candidate regions are restricted to those Networks. Resolution-group and exact-region candidates are then reranked using local discriminative genes in logCPM-compatible expression space. This hybrid route was selected because direct projected-VSD exact-region scoring was not consistently superior, especially in leave-one-monkey-out evaluation.

### 2.3 Interfaces and outputs

The Streamlit interface supports upload, metadata review, ranked Network and anatomical candidate display, downloadable tables and model warnings. The command-line interface supports scripted scoring, benchmark export and reproducible result bundles. Each report includes Top1 and Top3 candidates, rank scores, gene coverage, entropy, margins, route identifiers and resolution-specific labels. Accuracy is computed only when independently defined anatomical truth is available.

## 3 Implementation

cfRNA-BrainTrace is implemented in Python using NumPy, pandas and SciPy for scoring, scikit-learn utilities for evaluation, Plotly and Matplotlib for visualization, and Streamlit for the web interface. The Network projector route is implemented in `core/network_tracing.py`; the three-tier region route is implemented in `core/bo2023_region_tracing.py`; the web trigger and route text are exposed through `app/pages/tracing_page.py`. A smoke test on database sample `19R348` confirmed the production route `projected_vsd_network_top3_logcpm_resolution_local_exact`, Network overlap of 199/200 model genes and logCPM-based regional reranking.

## 4 Validation

### 4.1 Internal validation supports projected-VSD Network beam generation

In fold-local LOSO Network validation across 819 Bo2023 samples, projected VSD achieved Network Top1 accuracy of 0.5800 and Top3 accuracy of 0.9158, exceeding logCPM baseline (0.5556/0.8828) and native VSD (0.5311/0.8816). In strict LOMO validation, projected VSD achieved Network Top1 accuracy of 0.5372 and Top3 accuracy of 0.9133, again exceeding logCPM baseline and native VSD. These results support projected-VSD query representation for Network candidate generation.

### 4.2 Direct exact-region scoring is a diagnostic baseline, not the main endpoint

Direct exact-region projected-VSD scoring improved LOSO exact Top3 over logCPM baseline (0.3563 versus 0.2604), but in LOMO it remained below native VSD (0.3116 versus 0.3473). Exact-region Top1 and Top3 values were far below Network-level performance. These findings motivated the formal hybrid route and prevent claims that projected VSD is uniformly superior at exact-region resolution.

### 4.3 Formal three-tier hybrid validation

The complete three-tier route used projected-VSD Network Top3 beam generation followed by logCPM-compatible resolution-group and local exact-region reranking. In LOSO validation, Network, resolution-group and exact-region Top3 accuracies were 0.9238, 0.7236 and 0.4533, with median true ranks of 1.0, 2.0 and 4.0. In LOMO validation, the corresponding Top3 accuracies were 0.9121, 0.6909 and 0.4236. The gap between resolution-group and exact-region accuracy supports reporting resolution groups as the more defensible region-level endpoint.

### 4.4 External evaluations

In AHBA human brain RNA-seq with mapped labels, the hybrid route achieved Network Top1/Top3 accuracy of 0.7468/0.9442, group Top1/Top3 accuracy of 0.3626/0.6703 and exact Top1/Top3 accuracy of 0.2418/0.4286. Hybrid exact Top3 exceeded both logCPM baseline (0.3077) and projected-VSD-only scoring (0.2967), consistent with preserving Network-level projection gains while recovering local interpretability through logCPM reranking.

TCGA/BraTS glioma RNA-seq with corrected MRI-derived labels was used only for coarse anatomical consistency, because MRI truth labels are human atlas labels rather than Bo2023 macaque exact-region IDs. The hybrid route reached Network Top3 accuracy of 0.4000 and broad-anatomy Top3 accuracy of 0.6462, both above the logCPM baseline. GSE189919 projection established gene-overlap feasibility (15,622/21,668 projector genes; overlap fraction 0.7210) but was not used as a brain-origin accuracy validation because suitable anatomical truth labels were unavailable.

## 5 Conclusion

cfRNA-BrainTrace packages a three-tier atlas-based RNA source-tracing route into reproducible command-line and web interfaces. The evidence supports Network Top3 candidate generation and resolution-group reporting more strongly than exact-region Top1 localization. The software should therefore be used for hierarchical candidate ranking, quality control and transparent resolution auditing, not for direct clinical localization from unlabeled biofluid data.

## Funding

**[FUNDING INFORMATION TO BE INSERTED BEFORE SUBMISSION]**

## Conflict of Interest

None declared.

## Data availability

The Bo2023 macaque transcriptomic atlas, Allen Human Brain Atlas, TCGA/BraTS and GEO datasets are available from their original repositories under their original access conditions. Processed non-sensitive evaluation tables needed to reproduce manuscript figures should be included in the versioned software release.

## References

Bakas,S. *et al.* (2017) Advancing The Cancer Genome Atlas glioma MRI collections with expert segmentation labels and radiomic features. *Sci. Data*, **4**, 170117.

Bo,T. *et al.* (2023) Brain-wide and cell-specific transcriptomic insights into MRI-derived cortical morphology in macaque monkeys. *Nat. Commun.*, **14**, 1283.

Hawrylycz,M.J. *et al.* (2012) An anatomically comprehensive atlas of the adult human brain transcriptome. *Nature*, **489**, 391-399.

Vorperian,S.K. *et al.* (2022) Cell types of origin of the cell-free transcriptome. *Nat. Biotechnol.*, **40**, 855-861.

## Figure legend

**Fig. 1. Three-tier cfRNA-BrainTrace route and validation.** (A) Query expression is projected into Bo2023-like VSD space only for Network Top3 beam generation. (B) Resolution-group and exact-region candidates are reranked in local logCPM-compatible expression space. (C) Internal LOSO/LOMO validation supports Network Top3 and resolution-group reporting. (D) AHBA supports mapped-label external validation, whereas TCGA/BraTS supports only coarse anatomical consistency.

## Table

**Table 1. Principal route components and claim boundaries.**

| Component | Implementation | Manuscript claim boundary |
|---|---|---|
| Projector QC | Linear mapping from logCPM-compatible expression to Bo2023-like VSD space; sample-level global Pearson 0.9961 | Suitable for Network beam generation; not a new atlas and not gene-perfect VSD reconstruction |
| Network endpoint | Projected-VSD Top3 candidate beam | Main endpoint; LOSO Top3 0.9238 and LOMO Top3 0.9121 in formal route |
| Region-level endpoint | logCPM-compatible resolution-group reranking within Network beam | Primary region-level report; group Top3 exceeds exact Top3 in LOSO and LOMO |
| Exact-region output | Local logCPM exact-region reranking | Exploratory candidate list; do not present exact Top1 as deterministic localization |
| AHBA validation | Human brain RNA-seq with mapped labels | Cross-species mapped-label support; exact metrics only for evaluable labels |
| TCGA/BraTS validation | Glioma tissue RNA-seq plus MRI-derived human labels | Coarse anatomical consistency only; not Bo2023 exact-region validation |
"""


CN_MD = """# cfRNA-BrainTrace：基于灵长类转录组图谱的 RNA-seq 分层脑来源推断

**文章类型：** Bioinformatics Application Note  
**栏目：** Gene expression  
**作者与单位：** [投稿前补充]  
**通讯作者：** [投稿前补充]

## 摘要

### Summary

cfRNA-BrainTrace 是一个 Python 与 Streamlit 软件，用于从 RNA 表达谱中进行分层脑来源候选排序。当前锁定路线将 query 从 logCPM/logTPM 兼容表达空间投影到 Bo2023-like VSD 空间，仅用于 10 类 Network Top3 候选生成；随后 resolution group 与 exact region 均回到 logCPM 兼容的局部表达空间重排序。该设计避免把投影误写成新 atlas，并把解剖分辨率作为显式输出。在 Bo2023 内部验证中，正式三级 hybrid 路线的 Network Top3 在 LOSO 和 LOMO 中分别为 0.9238 和 0.9121；resolution-group Top3 分别为 0.7236 和 0.6909；exact-region Top3 分别为 0.4533 和 0.4236。AHBA 外部映射标签验证支持 hybrid 路线，TCGA/BraTS 只支持粗粒度解剖一致性。该软件适用于可复现的分层候选排序、质量控制和溯源分辨率边界审计。

### Availability and Implementation

Python 3.11+；支持命令行和 Streamlit 网页界面。开发仓库目前为私有仓库 **https://github.com/wz7717/cfrna-brain-tracing**；投稿前需完成公开发布、归档 DOI 和许可证确认。稳定在线应用地址为 **[投稿前补充]**。

## 1 引言

cfRNA 或组织 RNA 表达谱中可能保留来源组织信息，但脑内解剖定位受到区域转录相似性、图谱分辨率和样本域偏移限制。若直接给出单一 exact-region 标签，容易把粗粒度可信信号误写成精确定位。cfRNA-BrainTrace 因此采用分层报告策略：Network Top3 是主 endpoint，resolution group 是主要 region-level 粒度，exact region 仅作为局部候选排序。

## 2 软件与方法

### 2.1 输入与预处理

软件接收含基因符号和表达量的表格输入，并保留可用的样本、受试者、诊断和解剖元数据。输入被转换为 logCPM/logTPM 兼容空间后进行基因匹配、覆盖度审计和范围检查。每个样本都会输出匹配基因数、非零覆盖度、entropy、margin 和 route 标识。

### 2.2 Reference projection 与三级推断

训练参考仍为原始 Bo2023 表达资源：Network beam 使用 Bo2023 VSD/batch-removed 参考，resolution group 与 exact region 使用 raw-feature-count-derived logCPM 参考。projector 只把被留出样本或外部 query 映射到 Bo2023-like VSD 空间，用于 Network Top3 beam。projector 的 sample-level 拟合很强（global Pearson 0.9961，median sample Pearson 0.9969），但 gene-level 拟合较弱，因此不能写成“每个基因都被精确恢复”。

Network Top3 beam 生成后，候选脑区被限制在这些 Network 内；resolution group 与 exact region 再使用 logCPM 局部判别基因重排序。这个 hybrid 设计来自直接 exact-region 诊断结果：projected VSD 在 exact 层并非始终优于其他路线，尤其在 LOMO 中低于 native VSD。

## 3 实现

主线代码已经接入正式三级路线：Network projector route 位于 `core/network_tracing.py`，three-tier region route 位于 `core/bo2023_region_tracing.py`，网页触发和 route 说明位于 `app/pages/tracing_page.py`。数据库样本 `19R348` smoke test 确认输出 route 为 `projected_vsd_network_top3_logcpm_resolution_local_exact`，Network overlap 为 199/200，Region reference source 为 raw featurecounts logCPM。

## 4 验证结果

### 4.1 Network 层内部验证

在 819 个 Bo2023 样本的 fold-local LOSO Network 验证中，projected VSD 的 Network Top1/Top3 为 0.5800/0.9158，高于 logCPM baseline 和 native VSD。在 strict LOMO 验证中，projected VSD 的 Network Top1/Top3 为 0.5372/0.9133，同样高于两个对照路线。这是把 projected VSD 用作 Network candidate beam 的主要证据。

### 4.2 Direct exact-region 仅作为诊断基线

direct projected-VSD exact scoring 在 LOSO 中的 exact Top3 为 0.3563，高于 logCPM baseline；但在 LOMO 中 exact Top3 为 0.3116，低于 native VSD 的 0.3473。因此不能写“projected VSD 在 exact-region 层全面优于其他路线”，direct exact 也不应作为主线 endpoint。

### 4.3 正式三级 hybrid 验证

完整三级路线为 projected VSD Network Top3 beam -> logCPM resolution group rerank -> logCPM local exact rerank。LOSO 中 Network、resolution group、exact region 的 Top3 分别为 0.9238、0.7236、0.4533；LOMO 中分别为 0.9121、0.6909、0.4236。Group Top3 明显高于 Exact Top3，说明 resolution group 是更稳健的 region-level 报告粒度。

### 4.4 外部验证

AHBA human brain RNA-seq 映射标签验证中，hybrid 的 Network Top1/Top3 为 0.7468/0.9442，Group Top1/Top3 为 0.3626/0.6703，Exact Top1/Top3 为 0.2418/0.4286。hybrid 的 Exact Top3 高于 logCPM baseline 和 projected-VSD-only，支持“projected VSD 保留 Network 优势，logCPM downstream 修复 group/exact 层损失”的解释。

TCGA/BraTS glioma tissue RNA-seq + MRI label 只能用于 coarse anatomical consistency，因为 MRI 真值是人脑 atlas 标签，不是 Bo2023 macaque exact-region ID。hybrid 的 Network Top3 为 0.4000，Broad anatomy Top3 为 0.6462，均高于 logCPM baseline。GSE189919 只证明外部矩阵可投影到 projector gene space，不能写成 brain-origin accuracy 验证。

## 5 结论

cfRNA-BrainTrace 将三层 atlas-based RNA 溯源路线封装为可复现的软件。当前证据最支持 Network Top3 和 resolution-group 层级报告；exact-region Top1 仍受图谱分辨率和相邻脑区转录相似性限制。软件应用边界应定位为分层候选排序、质量控制和 claim-boundary 审计，而不是无标签 biofluid 的临床精确定位。

## 图注

**图1. cfRNA-BrainTrace 三级路线与验证。** query 仅在 Network beam 阶段投影到 Bo2023-like VSD；resolution group 和 exact region 回到 logCPM 局部空间重排序。内部 LOSO/LOMO 支持 Network Top3 与 group-level endpoint；AHBA 支持跨物种映射标签验证；TCGA/BraTS 仅支持粗粒度解剖一致性。
"""


SUPP_MD = """# Supplementary Material for cfRNA-BrainTrace

## Supplementary Methods

### S1 Reference atlas, gene-symbol audit and projector scope

The Bo2023 input resources were audited before manuscript use. Raw count, VSD and metadata matrices all contained 819 aligned samples. After symbol mapping, 21,668 common genes were available, and 199/200 locked Network-model genes were present in the common panel. Date-like Excel-mangled symbols were cleaned in the projector gene map and model artifact. The remaining suspicious identifiers were dominated by ENSMFAG fallback IDs or multi-ID aggregations and are reported as an interpretability limitation rather than a source-tracing failure.

### S2 Projector quality control

The projector was trained to map logCPM-compatible expression into Bo2023-like VSD space. Training-set sample-level reconstruction was strong (global Pearson 0.9961; median sample Pearson 0.9969; p10 sample Pearson 0.9940; MAE 0.2018; RMSE 0.2817). Gene-level agreement was weaker (median gene Pearson 0.6619; p10 gene Pearson 0.3665), so projected VSD was restricted to Network candidate generation.

### S3 Formal three-tier route

The production route first generates a Network Top3 beam in projected-VSD space. Candidate regions are then limited to those Networks. Resolution-group and exact-region candidates are reranked in logCPM-compatible local expression space using fold-local or route-local components. Accuracy is computed only when independent anatomical truth is available.

### S4 External validation policy

AHBA provides human normal brain RNA-seq with mapped labels and supports cross-species mapped-label validation. Exact-region metrics were calculated only for labels with stable Bo2023 mappings. TCGA/BraTS provides glioma tissue RNA-seq and corrected MRI-derived human labels; it supports only coarse anatomical consistency and is not a Bo2023 exact-region validation. GSE189919 was used only for projection feasibility because brain-origin accuracy labels were unavailable.

## Supplementary Results

### S1 Internal projected-VSD Network validation

In LOSO Network validation, projected VSD achieved Top1/Top3 accuracy of 0.5800/0.9158, compared with 0.5556/0.8828 for logCPM and 0.5311/0.8816 for native VSD. In LOMO validation, projected VSD achieved 0.5372/0.9133, compared with 0.4628/0.8510 for logCPM and 0.5043/0.8718 for native VSD.

### S2 Direct exact-region diagnostics

Direct exact-region scoring was weaker than Network scoring. In LOSO, projected-VSD exact Top1/Top3 was 0.1671/0.3563. In LOMO, projected-VSD exact Top1/Top3 was 0.1453/0.3116 and remained below native VSD Top3 of 0.3473. These diagnostics justify the hybrid downstream reranking route.

### S3 Formal three-tier validation

Formal LOSO hybrid validation yielded Network, group and exact Top3 accuracies of 0.9238, 0.7236 and 0.4533, respectively. Formal LOMO hybrid validation yielded 0.9121, 0.6909 and 0.4236. Exact-region performance remained clearly below Network performance, consistent with the intended interpretation hierarchy.

### S4 AHBA and special-label results

In AHBA mapped-label validation, the hybrid route achieved Network Top3 0.9442, Group Top3 0.6703 and Exact Top3 0.4286. Hybrid exact Top3 was higher than logCPM baseline (0.3077) and projected-VSD-only scoring (0.2967). For special labels, Insula, Caudate and Putamen all reached Network Top3 of 1.0000; Putamen exact Top3 was 0.8889, whereas Caudate exact Top1 was 0 and should not be described as an exact Top1 success.

### S5 TCGA/BraTS and GSE189919

In 65 TCGA/BraTS glioma tissue RNA-seq samples with MRI-derived labels, the hybrid route achieved Network Top3 0.4000 and Broad anatomy Top3 0.6462. These results support coarse anatomical consistency only. GSE189919 contained 51 samples and overlapped 15,622 of 21,668 projector genes (0.7210), supporting projection feasibility but not accuracy validation.

## Supplementary Tables

Table S1. Projector and gene-symbol audit.  
Table S2. Internal Network and formal three-tier validation.  
Table S3. Direct exact-region and local rerank diagnostics.  
Table S4. External AHBA, TCGA/BraTS and GSE189919 results.  
Table S5. Claim-boundary checklist for manuscript text.
"""


TABLES = {
    "Table1_route_components.csv": [
        ["Component", "Implementation", "Claim boundary"],
        ["Projector", "Query logCPM/logTPM-compatible expression -> Bo2023-like VSD", "Network beam only; not a new atlas"],
        ["Network", "Projected-VSD Top3 candidate generation", "Primary endpoint"],
        ["Resolution group", "logCPM local rerank within Network beam", "Main region-level endpoint"],
        ["Exact region", "logCPM local candidate rerank", "Exploratory local ranking"],
        ["External validation", "AHBA mapped labels; TCGA/BraTS coarse labels", "AHBA mapped-label support; TCGA/BraTS coarse consistency only"],
    ],
    "TableS1_projector_gene_audit.csv": [
        ["Metric", "Value", "Interpretation"],
        ["Common Bo2023 samples", "819", "Counts, VSD and metadata aligned"],
        ["Common genes after symbol mapping", "21668", "Projector gene space"],
        ["Locked Network genes in common panel", "199/200", "One prior date-like symbol issue cleaned"],
        ["Projector global Pearson", "0.9961", "Strong sample-level VSD-like reconstruction"],
        ["Projector median sample Pearson", "0.9969", "Supports Network beam use"],
        ["Projector median gene Pearson", "0.6619", "Do not claim gene-perfect reconstruction"],
        ["Date-like symbols after cleaning", "0", "Excel-mangled symbols removed"],
        ["Ensembl fallback symbols", "5499", "Interpretability limitation"],
    ],
    "TableS2_internal_formal_validation.csv": [
        ["Validation", "Level", "n", "Top1", "Top3", "Median true rank"],
        ["LOSO formal hybrid", "Network", "814", "0.5835", "0.9238", "1.0"],
        ["LOSO formal hybrid", "Resolution group", "814", "0.4447", "0.7236", "2.0"],
        ["LOSO formal hybrid", "Exact region", "814", "0.2248", "0.4533", "4.0"],
        ["LOMO formal hybrid", "Network", "819", "0.5775", "0.9121", "1.0"],
        ["LOMO formal hybrid", "Resolution group", "812", "0.4138", "0.6909", "2.0"],
        ["LOMO formal hybrid", "Exact region", "812", "0.2217", "0.4236", "5.0"],
    ],
    "TableS3_diagnostic_routes.csv": [
        ["Route", "Validation", "Level", "Top1", "Top3", "Interpretation"],
        ["Projected VSD direct", "LOSO", "Exact region", "0.1671", "0.3563", "Better than logCPM baseline but not enough as main endpoint"],
        ["Projected VSD direct", "LOMO", "Exact region", "0.1453", "0.3116", "Below native VSD Top3; diagnostic only"],
        ["Projected Network + logCPM local", "LOSO", "Exact region", "0.2088", "0.4263", "Supports downstream local rerank"],
        ["Projected Network + logCPM local", "Formal LOSO", "Exact region", "0.2248", "0.4533", "Best formal exact route but still exploratory"],
    ],
    "TableS4_external_validation.csv": [
        ["Dataset", "Route", "Endpoint", "Top1", "Top3", "Allowed claim"],
        ["AHBA", "Hybrid", "Network", "0.7468", "0.9442", "Cross-species mapped-label support"],
        ["AHBA", "Hybrid", "Resolution group", "0.3626", "0.6703", "Group-level mapped-label support"],
        ["AHBA", "Hybrid", "Exact region", "0.2418", "0.4286", "Only evaluable mapped labels"],
        ["TCGA/BraTS", "Hybrid", "Network", "0.1538", "0.4000", "Coarse anatomical consistency"],
        ["TCGA/BraTS", "Hybrid", "Broad anatomy", "0.1385", "0.6462", "Coarse anatomical consistency"],
        ["GSE189919", "Projection feasibility", "Gene overlap", "15622/21668", "0.7210", "No accuracy claim"],
    ],
    "TableS5_claim_boundaries.csv": [
        ["Do not write", "Use instead"],
        ["The whole Bo2023 atlas was reprojected into a new atlas.", "Only the query is projected to Bo2023-like VSD for Network beam generation."],
        ["Projected VSD is universally superior at exact-region resolution.", "Projected VSD supports Network beam generation; group/exact rerank uses logCPM local space."],
        ["TCGA/BraTS validates Bo2023 exact-region accuracy.", "TCGA/BraTS supports coarse anatomical consistency only."],
        ["Caudate exact Top1 performed well.", "Caudate/Putamen/Insula are stable at Network/group Top3; Caudate exact Top1 was 0."],
        ["Exact Top1 is a deterministic localization result.", "Exact outputs are local candidate rankings with clear limits."],
    ],
}


FIGURE_SOURCES = [
    (ROOT / "docs" / "figures" / "latest_three_tier_validation_summary.png", "Figure1_three_tier_validation_summary.png"),
    (RESULTS / "figures" / "projector_gene_qc_distributions.png", "FigureS1_projector_gene_qc_distributions.png"),
    (RESULTS / "figures" / "internal_network_accuracy.png", "FigureS2_internal_network_accuracy.png"),
    (RESULTS / "ahba_external_formal_three_tier" / "ahba_formal_three_tier_accuracy.png", "FigureS3_AHBA_formal_three_tier_accuracy.png"),
    (RESULTS / "tcga_labeled_hybrid_formal_external" / "tcga_labeled_hybrid_formal_accuracy.png", "FigureS4_TCGA_BraTS_coarse_consistency.png"),
]


def ensure_dirs() -> None:
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    TAB_OUT.mkdir(parents=True, exist_ok=True)


def write_texts() -> None:
    (OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_20260619.md").write_text(MAIN_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_CN_20260619.md").write_text(CN_MD, encoding="utf-8")
    (OUT / "Bioinformatics_Application_Note_Supplement_20260619.md").write_text(SUPP_MD, encoding="utf-8")
    for name, rows in TABLES.items():
        with (TAB_OUT / name).open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerows(rows)


def copy_figures() -> None:
    for src, name in FIGURE_SOURCES:
        if src.exists():
            shutil.copy2(src, FIG_OUT / name)


def set_font(run, size=10, bold=None, italic=None, name="Arial", east_asia="SimSun", color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_markdown(paragraph, text, size=10, east_asia="SimSun"):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size, east_asia=east_asia)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, east_asia=east_asia)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True, east_asia=east_asia)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, name="Consolas", east_asia=east_asia)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size, east_asia=east_asia)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc: Document, rows: list[list[str]], east_asia="SimSun") -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    usable = 6.5
    widths = [usable / len(rows[0])] * len(rows[0])
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.05
            set_font(para.add_run(value), size=8, bold=(i == 0), east_asia=east_asia)
            if i == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shading)
    if rows:
        repeat_header(table.rows[0])


def build_docx(md_path: Path, out_path: Path, cn=False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    east = "SimSun" if cn else "Arial"
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    normal.font.size = Pt(10.5 if cn else 10)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei" if cn else "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("BIOINFORMATICS | APPLICATION NOTE | 2026-06-19"), size=8, bold=True, color="6B7280", east_asia=east)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    table_rows: list[list[str]] = []
    in_table = False
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            if table_rows:
                add_table(doc, table_rows, east_asia=east)
                table_rows = []
            continue
        if line.startswith("|"):
            if "---" not in line:
                table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            in_table = True
            continue
        if in_table and table_rows:
            add_table(doc, table_rows, east_asia=east)
            table_rows = []
            in_table = False
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            add_markdown(p, line[2:], size=17 if cn else 16, east_asia=east)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(23, 54, 93)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_markdown(p, line[4:], size=12, east_asia=east)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_markdown(p, line[3:], size=13, east_asia=east)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_markdown(p, line, size=10.5 if cn else 10, east_asia=east)
    if table_rows:
        add_table(doc, table_rows, east_asia=east)

    fig = FIG_OUT / "Figure1_three_tier_validation_summary.png"
    if fig.exists():
        doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.add_run("Figure 1" if not cn else "图1")
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = pic.add_run().add_picture(str(fig), width=Inches(6.3))
        shape._inline.docPr.set("descr", "Three-tier cfRNA-BrainTrace validation summary")
        cap = doc.add_paragraph()
        add_markdown(cap, "Figure 1. Three-tier route and validation summary." if not cn else "图1. 三级路线与验证汇总。", size=8.5, east_asia=east)

    core = doc.core_properties
    core.title = md_path.stem
    core.subject = "Bioinformatics Application Note revision 2026-06-19"
    core.author = ""
    core.comments = "Generated from 2026-06-19 validation handoff guidance."
    doc.save(out_path)


def register_font() -> str:
    candidates = [r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\arial.ttf"]
    for candidate in candidates:
        if Path(candidate).exists():
            pdfmetrics.registerFont(TTFont("DocFont", candidate))
            return "DocFont"
    return "Helvetica"


def fit_image(path: Path, max_w: float, max_h: float) -> Image:
    with PILImage.open(path) as img:
        width, height = img.size
    scale = min(max_w / width, max_h / height)
    return Image(str(path), width=width * scale, height=height * scale)


def pdf_cell(text: str, style) -> Paragraph:
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def build_pdf_table(path: Path, content_w: float, body, header) -> Table:
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    data = [[pdf_cell(v, header) for v in rows[0]]]
    data.extend([[pdf_cell(v, body) for v in row] for row in rows[1:]])
    widths = [content_w / len(rows[0])] * len(rows[0])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FA")]),
    ]))
    return table


def build_supplement_pdf() -> None:
    font = register_font()
    page = landscape(A4)
    page_w, page_h = page
    margin = 12 * mm
    content_w = page_w - 2 * margin
    content_h = page_h - 2 * margin
    styles = getSampleStyleSheet()
    title = ParagraphStyle("Title", parent=styles["Title"], fontName=font, fontSize=19, leading=24, alignment=TA_CENTER, textColor=colors.HexColor("#17365D"))
    subtitle = ParagraphStyle("Subtitle", parent=styles["BodyText"], fontName=font, fontSize=10, leading=14, alignment=TA_CENTER)
    heading = ParagraphStyle("Heading", parent=styles["Heading1"], fontName=font, fontSize=13, leading=16, textColor=colors.HexColor("#1F4E79"), spaceAfter=5)
    caption = ParagraphStyle("Caption", parent=styles["BodyText"], fontName=font, fontSize=8.5, leading=11)
    body = ParagraphStyle("Cell", parent=styles["BodyText"], fontName=font, fontSize=7.2, leading=9)
    header = ParagraphStyle("Header", parent=body, fontName=font, textColor=colors.white, alignment=TA_CENTER)
    out_pdf = OUT / "Bioinformatics_Application_Note_Supplementary_Figures_Tables_20260619.pdf"
    doc = SimpleDocTemplate(str(out_pdf), pagesize=page, leftMargin=margin, rightMargin=margin, topMargin=margin, bottomMargin=margin)
    story = [
        Spacer(1, 30 * mm),
        Paragraph("cfRNA-BrainTrace Supplementary Figures and Tables", title),
        Paragraph("Updated for the 2026-06-19 three-tier hybrid validation route", subtitle),
        Spacer(1, 8 * mm),
        Paragraph("Includes copied validation figures and regenerated supplementary tables.", subtitle),
        PageBreak(),
    ]
    captions = {
        "Figure1_three_tier_validation_summary.png": "Figure 1. Three-tier validation summary for projected-VSD Network beam and logCPM downstream reranking.",
        "FigureS1_projector_gene_qc_distributions.png": "Figure S1. Projector gene-level and sample-level QC distributions.",
        "FigureS2_internal_network_accuracy.png": "Figure S2. Internal Network accuracy comparing logCPM, native VSD and projected VSD routes.",
        "FigureS3_AHBA_formal_three_tier_accuracy.png": "Figure S3. AHBA formal three-tier mapped-label external validation.",
        "FigureS4_TCGA_BraTS_coarse_consistency.png": "Figure S4. TCGA/BraTS glioma RNA-seq and MRI-label coarse anatomical consistency.",
    }
    for fig in sorted(FIG_OUT.glob("*.png")):
        story.append(Paragraph(fig.stem.replace("_", " "), heading))
        story.append(fit_image(fig, content_w, content_h - 28 * mm))
        story.append(Paragraph(captions.get(fig.name, fig.name), caption))
        story.append(PageBreak())
    for idx, table_path in enumerate(sorted(TAB_OUT.glob("*.csv"))):
        story.append(Paragraph(table_path.stem.replace("_", " "), heading))
        story.append(build_pdf_table(table_path, content_w, body, header))
        if idx < len(list(TAB_OUT.glob("*.csv"))) - 1:
            story.append(PageBreak())
    doc.build(story)


def main() -> None:
    ensure_dirs()
    write_texts()
    copy_figures()
    build_docx(OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_20260619.md", OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_20260619.docx")
    build_docx(OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_CN_20260619.md", OUT / "Bioinformatics_Application_Note_cfRNA_BrainTrace_CN_20260619.docx", cn=True)
    build_docx(OUT / "Bioinformatics_Application_Note_Supplement_20260619.md", OUT / "Bioinformatics_Application_Note_Supplement_20260619.docx")
    build_supplement_pdf()
    print(OUT)


if __name__ == "__main__":
    main()
