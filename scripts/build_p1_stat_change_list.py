#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

out = Path("manuscript/BrainTrace_v5_round4_P1-STAT1-4_修改清单.docx")
d = Document()
sec = d.sections[0]
sec.top_margin = sec.bottom_margin = Inches(.72)
sec.left_margin = sec.right_margin = Inches(.75)
styles = d.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9.5)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(18)

p = d.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("BrainTrace v5 第四轮：统计方法学 P1（STAT1-4）修改清单")
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("基线：BIOP1-4 最新累积稿；输出：STATP1-4 最新累积稿；Word 修订模式保留").italic = True

d.add_heading("一、审阅结论", level=1)
d.add_paragraph(
    "四项均已按审稿建议完成方法评估、可复现计算和同步修订。STAT1 改变多重检验家族后仍无结果通过 "
    "BH 0.05；STAT2 明确 MDE 不以样本级 ICC 为输入，并补充设计效应/方差分解敏感性；STAT3 "
    "新增同情景下一次扰动重复的 95% 预测区间；STAT4 将 n=9 下的稳健性措辞降级为方向一致、功效有限。"
)

rows = [
    ("P1-STAT1", "BH 家族遗漏 Network Top1",
     "无外部时间戳预注册证据可支持排除，故采用四检验家族。",
     "raw P/BH q：Network Top1 0.031250/0.125000；Network Top3 0.375000/0.500000；"
     "Group Top3 0.593750/0.593750；Exact Top3 0.324219/0.500000。均未通过 0.05。",
     "主文 §4；补充 S3、S11、S13、Table S16"),
    ("P1-STAT2", "MDE 未说明 ICC 和方差分解",
     "MDE 单位为 9 个供体配对率差，ICC 不是模拟输入；另做假设索引敏感性，避免伪称经验 ICC。",
     "平均簇大小 91；ICC=0/.01/.05/.10/.30 时设计效应=1/1.9/5.5/10/28，"
     "有效样本量=819.0/431.1/148.9/81.9/29.3。p=0.9176 时总 Bernoulli 方差=0.07561，"
     "补充材料列出供体间/供体内分量。",
     "主文 §4 新增审计段；补充 S3 新增审计段"),
    ("P1-STAT3", "稀疏查询仅报告 CI",
     "按 30 次重复的均值和 SD，计算下一次同机制扰动重复的正态理论预测区间："
     "mean ± t(0.975,29)×SD×sqrt(1+1/30)。",
     "12 个区间已写入 Table S15。示例：Extreme Network 55.73-61.35%；"
     "Group 36.31-43.37%；Exact 14.09-19.67%。明确不是临床单查询或供体总体 PI。",
     "主文 §5；补充 S15、Table S15 notes"),
    ("P1-STAT4", "n=9 下“robust”表述过强",
     "统一改为 directionally consistent / limited power，并明确非显著不等于等效、无差异或全局稳健。",
     "主文 donor-macro、BH/MDE 与 domain-shift 解释均已降级；S19 超参数说明同步修改。",
     "主文 §4、§5、Figure 1 图注；补充 S13、S19"),
]
t = d.add_table(rows=1, cols=5)
t.style = "Table Grid"
headers = ["编号", "问题", "修复方法", "计算/结论", "落点"]
for c, h in zip(t.rows[0].cells, headers):
    c.text = h
for row in rows:
    cells = t.add_row().cells
    for c, val in zip(cells, row):
        c.text = val
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
widths = [0.75, 1.2, 1.75, 2.65, 1.25]
for row in t.rows:
    for c, w in zip(row.cells, widths):
        c.width = Inches(w)

d.add_heading("二、预测区间核对表（%）", level=1)
pi = [
    ("Mild", "90.13-92.09", "69.23-72.69", "40.07-43.97"),
    ("Moderate", "81.91-85.77", "60.56-65.72", "30.35-36.13"),
    ("Severe", "69.42-74.24", "49.84-54.08", "22.07-26.93"),
    ("Extreme", "55.73-61.35", "36.31-43.37", "14.09-19.67"),
]
t2 = d.add_table(rows=1, cols=4)
t2.style = "Table Grid"
for c, h in zip(t2.rows[0].cells, ["情景", "Network Top3", "Resolution-group Top3", "Exact-region Top3"]):
    c.text = h
for r in pi:
    for c, v in zip(t2.add_row().cells, r):
        c.text = v

d.add_heading("三、同步检查与质量门", level=1)
checks = [
    "主文：§4 多重检验/MDE、§5 稀疏输入与 domain-shift 解释已同步。",
    "图注：Figure 1 已标明内部结果来自 9 个供体，不能建立稳健性。",
    "补充材料：S3、S11、S13、S15、S19 及 Tables S15-S16 已同步。",
    "计算产物：manuscript/calculations/P1_STAT1-4_audit.json 与 P1_STAT1-4_prediction_intervals.csv。",
    "结构 QA：trackRevisions 开启；修订节点无嵌套；BH 数值、12 个 PI 与正文/表格一致。",
    "版式 QA：主文、补充材料及本清单均逐页渲染检查。"
]
for x in checks:
    d.add_paragraph(x, style="List Bullet")

d.add_heading("四、审阅时建议重点", level=1)
d.add_paragraph(
    "请重点确认：(1) 是否接受在缺乏可核验预注册记录时将 Network Top1 纳入四检验家族；"
    "(2) MDE 的供体级模拟解释是否充分；(3) Table S15 的 repeat PI 标签不会被误读为单患者/单查询预测区间；"
    "(4) 全文结论强度已与 n=9 的信息量相匹配。"
)
out.parent.mkdir(parents=True, exist_ok=True)
d.save(out)
print(out)
