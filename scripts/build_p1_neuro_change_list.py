#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

out = Path("manuscript/BrainTrace_v5_round4_P1-NEURO1-4_修改清单.docx")
d = Document()
s = d.sections[0]
s.top_margin = s.bottom_margin = Inches(.7)
s.left_margin = s.right_margin = Inches(.72)
d.styles["Normal"].font.name = "Arial"
d.styles["Normal"].font.size = Pt(9.5)
d.styles["Title"].font.name = "Arial"
d.styles["Title"].font.size = Pt(18)
p = d.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("BrainTrace v5 第四轮：神经科学 P1（NEURO1-4）修改清单")
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("基线：STATP1-4 最新累积稿；输出：NEUROP1-4 最新累积稿；Word 修订模式保留").italic = True

d.add_heading("一、总体结论", level=1)
d.add_paragraph(
    "四项问题均完成定量审计并同步修订。23个 resolution groups 中17个具有明确的局部/系统内解剖解释，"
    "6个必须限定为跨亚区的操作性转录组分组；Hippocampal 单区域退化已落实为工具输出警告；"
    "Temporal 与 Parietal 的混淆已按冻结 LOSO 矩阵分解，并将神经科学解释限定为与多模态汇聚相容、非机制证明。"
)

items = [
    ("P1-NEURO1", "23个 resolution groups 缺乏系统解剖验证",
     "用两级审计规则逐组评估：同一局部场族/功能解剖系统=clear；跨细胞构筑或感觉-联合亚区=qualified。",
     "17/23 clear，6/23 qualified。六组为：10o/46d/46v；12l/12r/45/8A；G/Id/Ig/SII；"
     "13a/13b/25；5/LIPv/VIP；CL/ML/Tpt。",
     "主文 §2 新增审计段；补充 Table S17 及其后审计段；逐行 CSV 已归档。"),
    ("P1-NEURO2", "Hippocampal 三层结构退化讨论及工具标记不足",
     "将其定义为结构性退化而非三级独立证据，并修改工具使警告不依赖可选 resolution model。",
     "1个区域、8个训练样本；输出 single_region_network=true、tier_degeneracy='group=exact'。"
     "3项注释测试通过，包括缺失模型路径。",
     "主文 §5、Figure 1 图注/alt text；补充 S11、Supplementary Figure S1。"),
    ("P1-NEURO3", "Temporal 内部异质性对混淆贡献未量化",
     "从冻结 LOSO 混淆矩阵分解 Temporal 的错误目的地。",
     "193个真值样本：85正确、108假阴性；47到 Operculum/Insula、46到 Parietal，合计86.1%；"
     "再加13到 Visual/dorsal STS 后为98.1%。",
     "主文 §2；补充 S13。"),
    ("P1-NEURO4", "Parietal 低精度/高召回缺乏多模态汇聚解释",
     "由 recall/precision 与冻结矩阵重建 TP、预测总数和主要假阳性来源。",
     "TP=78/98，recall=0.7959；约177个 Parietal 预测、约99个假阳性；Temporal 46加"
     "Visual/dorsal STS 13至少占59/99=59.6%。",
     "主文 §2；补充 S13；明确未检验连接或因果机制。"),
]
t = d.add_table(rows=1, cols=5)
t.style = "Table Grid"
for c, h in zip(t.rows[0].cells, ["编号", "审稿问题", "修复方法", "计算结果", "同步位置"]):
    c.text = h
for item in items:
    for c, value in zip(t.add_row().cells, item):
        c.text = value

d.add_heading("二、6个 qualified resolution groups 的审阅边界", level=1)
for text in [
    "LPFC 10o/46d/46v：额极、背外侧和腹外侧前额叶混合。",
    "LPFC 12l/12r/45/8A：眶部/腹外侧与背侧 area-8 系统混合。",
    "Operculum/Insula G/Id/Ig/SII：岛叶亚区与第二躯体感觉盖区合并。",
    "OMPFC 13a/13b/25：眶额 area 13 与膝下 area 25 合并。",
    "Parietal 5/LIPv/VIP：躯体感觉联合区与顶内沟视觉空间区合并。",
    "Temporal CL/ML/Tpt：听觉带区与颞顶联合区合并。"
]:
    d.add_paragraph(text, style="List Bullet")
d.add_paragraph(
    "上述分组可作为锁定算法的操作性不确定集合使用，但名称不得被解释为经典细胞构筑单位或已证实的功能连接模块。"
)

d.add_heading("三、质量检查", level=1)
for text in [
    "计算文件：P1_NEURO1-4_audit.json；逐组审计：P1_NEURO1_resolution_group_anatomy_audit.csv。",
    "代码：Hippocampal 特例在可选模型查找前执行；不改变分数或排序。",
    "图注：Figure 1、alt text 与 Supplementary Figure S1 已同步 single-region/group=exact 警告。",
    "补充材料：S11、S13、Table S17 已同步。",
    "结构 QA：trackRevisions 开启，无嵌套修订；计算值与正文/补充材料一致。",
    "版式 QA：主文、补充材料和本清单逐页渲染检查。"
]:
    d.add_paragraph(text, style="List Bullet")

d.add_heading("四、建议审阅重点", level=1)
d.add_paragraph(
    "请重点确认：(1) 两级解剖审计规则是否足以支撑“17 clear / 6 qualified”；"
    "(2) qualified 分组是否应继续保留中性操作性名称；(3) Hippocampal 的重复细层输出不会被误读为独立证据；"
    "(4) Temporal/Parietal 的多模态汇聚解释已明确限定为与数据相容而非机制证明。"
)
out.parent.mkdir(parents=True, exist_ok=True)
d.save(out)
print(out)
