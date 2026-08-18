#!/usr/bin/env python
"""Apply minimal final scientific-remediation text changes to submission DOCX files.

All numerical replacements come from SCIENTIFIC_REMEDIATION_QA.json and the
staged benchmark manifest.  This script never modifies the source documents.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DIR = ROOT / "manuscript_remediation"
DEFAULT_MAIN_IN = DEFAULT_DIR / "BrainTrace_Main_Manuscript_HuangRemediated.docx"
DEFAULT_SUPP_IN = DEFAULT_DIR / "BrainTrace_Supplementary_File_HuangRemediated.docx"
DEFAULT_MAIN_OUT = DEFAULT_DIR / "BrainTrace_Main_Manuscript_FinalScientificRemediated.docx"
DEFAULT_SUPP_OUT = DEFAULT_DIR / "BrainTrace_Supplementary_File_FinalScientificRemediated.docx"
QA_PATH = ROOT / "SCIENTIFIC_REMEDIATION_QA.json"
BENCHMARK_PATH = ROOT / "reproducibility" / "formal_real_input_performance_manifest.json"


def replace_once_in_paragraph(paragraph, old: str, new: str, label: str) -> None:
    """Replace one text span while preserving surrounding run formatting."""

    text = paragraph.text
    count = text.count(old)
    if count != 1:
        raise ValueError(f"{label}: expected one occurrence, found {count}")
    start = text.index(old)
    end = start + len(old)
    cursor = 0
    start_run = end_run = None
    for index, run in enumerate(paragraph.runs):
        next_cursor = cursor + len(run.text)
        if start_run is None and cursor <= start < next_cursor:
            start_run = (index, start - cursor)
        if cursor < end <= next_cursor:
            end_run = (index, end - cursor)
            break
        cursor = next_cursor
    if start_run is None or end_run is None:
        raise ValueError(f"{label}: could not locate text span in runs")
    first_index, first_offset = start_run
    last_index, last_offset = end_run
    if first_index == last_index:
        run = paragraph.runs[first_index]
        run.text = run.text[:first_offset] + new + run.text[last_offset:]
        return
    first = paragraph.runs[first_index]
    first.text = first.text[:first_offset] + new
    for index in range(first_index + 1, last_index):
        paragraph.runs[index].text = ""
    last = paragraph.runs[last_index]
    last.text = last.text[last_offset:]


def replace_at(document: Document, paragraph_index: int, old: str, new: str, label: str) -> None:
    replace_once_in_paragraph(document.paragraphs[paragraph_index], old, new, label)


def replace_table_cell(table, row: int, column: int, old: str, new: str, label: str) -> None:
    cell = table.cell(row, column)
    if len(cell.paragraphs) != 1:
        raise ValueError(f"{label}: expected one cell paragraph")
    replace_once_in_paragraph(cell.paragraphs[0], old, new, label)


def all_document_text(document: Document) -> str:
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        blocks.extend(paragraph.text for paragraph in section.header.paragraphs)
        blocks.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(blocks)


def update_main(document: Document, qa: dict) -> None:
    lomo = qa["lomo_exact_f1"]["summary"]
    benchmark = qa["benchmark"]
    replace_at(
        document,
        28,
        "On 51 GSE189919 profiles, cold inference required 0.384 s/sample with a 222 MiB peak working set;",
        (
            f"The benchmark input comprised {benchmark['n_profiles']} profiles × {benchmark['n_genes']:,} genes. "
            f"Cold frozen-route inference required {benchmark['cold']['total_seconds']:.4f} s in total "
            f"({benchmark['cold']['seconds_per_profile']:.4f} s/profile) and peaked at "
            f"{benchmark['cold']['peak_working_set_mib']:.1f} MiB; across {benchmark['n_warm_repeats']} warm repeats "
            f"({benchmark['warm_events']} timed inference events), the maximum observed working set was "
            f"{benchmark['warm']['maximum_working_set_mib']:.1f} MiB;"
        ),
        "main benchmark",
    )
    replace_at(
        document,
        30,
        "Exact-region macro-F1 (mean\u00B1SD) was 0.210\u00B10.225 (105 classes) in LOSO and 0.194\u00B10.193 (104 classes) in LOMO; medians (IQR) were 0.154 (0.286) and 0.132 (0.179).",
        (
            "Exact-region macro-F1 (mean±SD) was 0.210±0.225 (105 classes) in LOSO and "
            f"{lomo['macro_f1']:.3f}±{lomo['sd_class_f1']:.3f} (104 classes) in LOMO; medians (IQR) were "
            f"0.154 (0.286) and {lomo['median_class_f1']:.3f} ({lomo['iqr_class_f1']:.3f})."
        ),
        "main LOMO Exact F1",
    )


def update_supplement(document: Document, qa: dict, benchmark_manifest: dict) -> None:
    tcga = qa["tcga_broad_strict_top3_range"]
    lomo = qa["lomo_exact_f1"]["summary"]
    benchmark = qa["benchmark"]
    friedman = qa["friedman_exact_enumeration"]
    baselines = {
        row["endpoint"]: row for row in qa["resolution_group_random_baselines"]["records"]
    }
    cold_timing = benchmark_manifest["cold"]["timing"]
    warm_timing = benchmark_manifest["warm"]["aggregate"]
    warm_python_peak_mib = max(
        repeat["memory"]["python_peak_bytes"] / 1024**2
        for repeat in benchmark_manifest["warm"]["repeats"]
    )

    replace_at(
        document,
        54,
        "the formal LOMO Network prediction-level source formal_lomo_network_detail.csv, integer-count class metrics formal_lomo_network_f1.csv and lomo_network_f1_provenance.json document the regenerated F1 evidence chain.",
        "the formal LOMO Network prediction-level source formal_lomo_network_detail.csv, integer-count class metrics formal_lomo_network_f1.csv and lomo_network_f1_provenance.json, and the formal LOMO Exact prediction-level source formal_lomo_exact_region_detail.csv, integer-count class metrics formal_lomo_exact_region_f1.csv and lomo_exact_region_f1_provenance.json document the regenerated F1 evidence chains.",
        "supplement F1 provenance index",
    )
    replace_at(
        document,
        61,
        "In LOMO, the corresponding values were 0.194\u00B10.193 (104 classes), 0.132 (IQR 0.179), 0.210 and 0.212; 12/104 classes had F1=0 and conditional macro-F1 was 0.220.",
        (
            f"In LOMO, the corresponding values were {lomo['macro_f1']:.3f}±{lomo['sd_class_f1']:.3f} (104 classes), "
            f"{lomo['median_class_f1']:.3f} (IQR {lomo['iqr_class_f1']:.3f}), "
            f"{lomo['weighted_f1']:.3f} and {lomo['micro_f1']:.3f}; "
            f"{lomo['n_zero_f1_classes']}/104 classes had F1=0 and conditional macro-F1 was "
            f"{lomo['conditional_macro_f1_nonzero']:.3f}."
        ),
        "supplement R1 LOMO Exact F1",
    )
    replace_at(
        document,
        70,
        "range 32.02 points",
        f"range {tcga['derived_range_percentage_points']:.2f} percentage points",
        "supplement TCGA broad range",
    )
    replace_at(
        document,
        80,
        "Friedman hit3 chi2=0.5385, df=2, with both the asymptotic and 19,683-pattern exact-enumeration P=0.764; hit1 P=0.223.",
        f"Friedman hit3 chi2={friedman['chi2']:.4f}, df={friedman['df']}, P={friedman['p_value']:.3f}; hit1 P=0.223.",
        "supplement Friedman sensitivity",
    )
    replace_table_cell(
        document.tables[9],
        16,
        2,
        "Given n=9, the chi-squared approximation may deviate slightly from the exact Friedman distribution.",
        "Given n=9, this chi-squared approximation is the only reported Friedman inference.",
        "Table S15 Friedman scope",
    )
    replace_at(
        document,
        82,
        "Formal real-input engineering performance used all 51 public GSE189919 raw-count samples in fixed matrix-header order at concurrency 1. Cold process-entry-to-completion time was 20.9448 s; cold frozen-route inference was 19.5901 s (0.3841 s/sample; p50/p95 0.3769/0.4026 s). Three warm repeats totalled 59.6211 s for 153 timed inferences (0.3897 s/sample; pooled p50/p95 0.3797/0.4252 s). Peak working set was 222.0 MiB and peak tracemalloc allocation was 76.7 MiB. ",
        (
            f"Formal real-input engineering performance used a {benchmark['n_profiles']}-profile × {benchmark['n_genes']:,}-gene public GSE189919 raw-count matrix in fixed header order at concurrency 1. "
            f"Cold process-entry-to-completion time was {benchmark_manifest['cold']['process_entry_to_completion_seconds']:.4f} s; cold frozen-route inference was "
            f"{benchmark['cold']['total_seconds']:.4f} s ({benchmark['cold']['seconds_per_profile']:.4f} s/profile; "
            f"p50/p95 {cold_timing['sample_time_p50_seconds']:.4f}/{cold_timing['sample_time_p95_seconds']:.4f} s) and cold peak working set was "
            f"{benchmark['cold']['peak_working_set_mib']:.1f} MiB. Across {benchmark['n_warm_repeats']} warm repeats "
            f"({benchmark['warm_events']} timed inference events), total time was {benchmark['warm']['total_seconds']:.4f} s "
            f"({benchmark['warm']['seconds_per_event']:.4f} s/event; pooled p50/p95 "
            f"{warm_timing['sample_time_p50_seconds']:.4f}/{warm_timing['sample_time_p95_seconds']:.4f} s), maximum observed working set was "
            f"{benchmark['warm']['maximum_working_set_mib']:.1f} MiB, and peak tracemalloc allocation was {warm_python_peak_mib:.1f} MiB. "
        ),
        "supplement benchmark stages",
    )
    replace_at(
        document,
        103,
        "Exact-region class F1 was LOSO mean\u00B1SD 0.210\u00B10.225 (105 classes), median 0.1538, Q1 0, Q3 0.2857, IQR 0.2857, and LOMO mean\u00B1SD 0.194\u00B10.193 (104 classes), median 0.13245, Q1 0.06648, Q3 0.24543, IQR 0.17895.",
        (
            "Exact-region class F1 was LOSO mean±SD 0.210±0.225 (105 classes), median 0.1538, Q1 0, Q3 0.2857, IQR 0.2857, and "
            f"LOMO mean±SD {lomo['macro_f1']:.3f}±{lomo['sd_class_f1']:.3f} (104 classes), "
            f"median {lomo['median_class_f1']:.5f}, Q1 {lomo['q1_class_f1']:.0f}, "
            f"Q3 {lomo['q3_class_f1']:.5f}, IQR {lomo['iqr_class_f1']:.5f}."
        ),
        "supplement R9 LOMO Exact F1",
    )
    replace_at(
        document,
        103,
        "Friedman results report both the chi-square approximation and exact enumeration of 19,683 rank patterns because only nine matched donors were observed.",
        f"Friedman results use the chi-square approximation (χ²={friedman['chi2']:.4f}, df={friedman['df']}, P={friedman['p_value']:.3f}) across nine matched donors.",
        "supplement R9 Friedman",
    )
    replace_at(
        document,
        117,
        "The benchmarked 153-sample by 28,415-gene raw-count input peaked at 222.0 MiB.",
        (
            f"The benchmark used a {benchmark['n_profiles']}-profile × {benchmark['n_genes']:,}-gene raw-count input; across "
            f"{benchmark['warm_events']} warm timed inference events, the maximum observed working set was "
            f"{benchmark['warm']['maximum_working_set_mib']:.1f} MiB."
        ),
        "supplement benchmark unit",
    )
    table = document.tables[8]
    replace_table_cell(
        table, 3, 1, "22.0%", f"{baselines['LOSO']['uniform_random_rate'] * 100:.1f}%", "Table S8 LOSO uniform"
    )
    replace_table_cell(
        table, 3, 2, "6.3%", f"{baselines['LOSO']['weighted_random_rate'] * 100:.1f}%", "Table S8 LOSO weighted"
    )
    replace_table_cell(
        table, 4, 1, "20.4%", f"{baselines['LOMO']['uniform_random_rate'] * 100:.1f}%", "Table S8 LOMO uniform"
    )
    replace_table_cell(
        table, 4, 2, "4.4%", f"{baselines['LOMO']['weighted_random_rate'] * 100:.1f}%", "Table S8 LOMO weighted"
    )


def assert_document_content(document: Document, qa: dict) -> None:
    text = all_document_text(document)
    forbidden = (
        "range 32.02 points",
        "153-sample by 28,415-gene",
        "19,683-pattern exact-enumeration",
        "exact enumeration of 19,683",
        "exact Friedman distribution",
        "0.194\u00B10.193 (104 classes)",
        "12/104 classes had F1=0",
    )
    present = [value for value in forbidden if value in text]
    if present:
        raise ValueError(f"Stale scientific text remains: {present}")
    lomo = qa["lomo_exact_f1"]["summary"]
    if f"{lomo['macro_f1']:.3f}" not in text:
        raise ValueError("Current LOMO Exact macro-F1 is absent from updated document")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--main-input", type=Path, default=DEFAULT_MAIN_IN)
    parser.add_argument("--supplement-input", type=Path, default=DEFAULT_SUPP_IN)
    parser.add_argument("--main-output", type=Path, default=DEFAULT_MAIN_OUT)
    parser.add_argument("--supplement-output", type=Path, default=DEFAULT_SUPP_OUT)
    parser.add_argument("--qa", type=Path, default=QA_PATH)
    args = parser.parse_args()

    if args.main_input.resolve() == args.main_output.resolve():
        raise ValueError("Main output must not overwrite the source document")
    if args.supplement_input.resolve() == args.supplement_output.resolve():
        raise ValueError("Supplement output must not overwrite the source document")
    qa = json.loads(args.qa.read_text(encoding="utf-8"))
    benchmark_manifest = json.loads(BENCHMARK_PATH.read_text(encoding="utf-8"))

    main_document = Document(args.main_input)
    supplement_document = Document(args.supplement_input)
    update_main(main_document, qa)
    update_supplement(supplement_document, qa, benchmark_manifest)
    assert_document_content(main_document, qa)
    assert_document_content(supplement_document, qa)
    args.main_output.parent.mkdir(parents=True, exist_ok=True)
    main_document.save(args.main_output)
    supplement_document.save(args.supplement_output)
    print(args.main_output)
    print(args.supplement_output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
