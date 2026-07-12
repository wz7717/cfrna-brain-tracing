"""Clarify primary cluster-robust BH inference versus McNemar sensitivity analysis."""

from copy import deepcopy
from pathlib import Path
import sys
import zipfile

from docx import Document


ROOT = Path(r"D:\Download\cfrna-brain-tracing-streamlit-cloud-ready")
MAIN_SOURCE = ROOT / "0711Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_P0_review_metrics_figure1_updated.docx"
MAIN_OUTPUT = ROOT / "0711Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_P0_review_metrics_figure1_cluster_bh_clarified.docx"
SUPP_SOURCE = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_complete_with_sparse_donor_tables.docx"
SUPP_OUTPUT = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_complete_with_sparse_donor_tables_cluster_bh_clarified.docx"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for run in paragraph.runs:
        paragraph._p.remove(run._r)
    run = paragraph.add_run(new_text)
    if rpr is not None:
        run._r.insert(0, rpr)


def find_paragraph(doc, prefix: str):
    matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def image_bytes(path: Path) -> bytes:
    with zipfile.ZipFile(path) as package:
        return package.read("word/media/image1.png")


def update_main() -> None:
    source_image = image_bytes(MAIN_SOURCE)
    doc = Document(MAIN_SOURCE)
    paragraph = find_paragraph(doc, "The complete validation tested the three-tier architecture")
    old = paragraph.text
    old_fragment = (
        "After Benjamini-Hochberg correction across the three prespecified hierarchical Top3 comparisons, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold. "
        "The original paired exact McNemar p values were 0.1686, 0.0389 and 0.0532 for Network, resolution group and exact region, respectively; the corresponding BH-adjusted values were 0.1686, 0.0798 and 0.0798."
    )
    new_fragment = (
        "The corresponding cluster-robust BH-adjusted p values were 0.0688, 0.0828 and 0.1959; thus, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold after correction. "
        "As a paired sample-level sensitivity analysis, exact McNemar p values were 0.1686, 0.0389 and 0.0532 for Network, resolution group and exact region, respectively; its separately BH-adjusted values were 0.1686, 0.0798 and 0.0798."
    )
    if old_fragment not in old:
        raise RuntimeError("Expected main-text BH/McNemar fragment was not found")
    replace_paragraph_text(paragraph, old.replace(old_fragment, new_fragment))
    doc.save(MAIN_OUTPUT)
    if image_bytes(MAIN_OUTPUT) != source_image:
        raise RuntimeError("Embedded Figure 1 changed unexpectedly")


def update_supplement() -> None:
    doc = Document(SUPP_SOURCE)
    results_paragraph = find_paragraph(doc, "Wilson confidence intervals (Wilson, 1927)")
    old = results_paragraph.text
    old_fragment = (
        "After Benjamini-Hochberg correction across the three prespecified hierarchical Top3 comparisons, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold. "
        "The original paired exact McNemar p values were 0.1686, 0.0389 and 0.0532 for Network, resolution group and exact region, respectively; the corresponding BH-adjusted values were 0.1686, 0.0798 and 0.0798."
    )
    new_fragment = (
        "The corresponding cluster-robust BH-adjusted p values were 0.0688, 0.0828 and 0.1959; thus, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold after correction. "
        "As a paired sample-level sensitivity analysis, exact McNemar p values were 0.1686, 0.0389 and 0.0532 for Network, resolution group and exact region, respectively; its separately BH-adjusted values were 0.1686, 0.0798 and 0.0798."
    )
    if old_fragment not in old:
        raise RuntimeError("Expected supplementary results BH/McNemar fragment was not found")
    replace_paragraph_text(results_paragraph, old.replace(old_fragment, new_fragment))

    clarification = find_paragraph(doc, "The three prespecified internal hierarchical Top3 comparisons")
    replace_paragraph_text(
        clarification,
        "For the primary donor-clustered inference, the three prespecified internal hierarchical Top3 comparisons were corrected for multiple testing using the Benjamini-Hochberg procedure. "
        "The cluster-robust adjusted p values were 0.0688 for Network, 0.0828 for resolution group and 0.1959 for exact region; thus, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold after correction. "
        "Paired exact McNemar tests were retained as sample-level sensitivity analyses: their separately BH-adjusted p values were 0.1686 for Network and 0.0798 for both resolution group and exact region. "
        "Endpoint-specific random-baseline tests and coarse external-consistency analyses address separate descriptive questions and are reported with their stated interpretation boundaries."
    )

    table_s13 = doc.tables[-1]
    header_cell = table_s13.rows[0].cells[5]
    replace_paragraph_text(header_cell.paragraphs[0], "Cluster-robust\np raw / BH")
    doc.save(SUPP_OUTPUT)


def main() -> None:
    update_main()
    update_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
    print("Clarified cluster-robust primary inference and McNemar sensitivity analysis")


if __name__ == "__main__":
    main()
