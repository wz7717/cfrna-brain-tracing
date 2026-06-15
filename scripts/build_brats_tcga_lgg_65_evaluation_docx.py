#!/usr/bin/env python
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / "results" / "brats_tcga_lgg_65_mri_truth_evaluation_20260609"
OUT = OUTDIR / "TCGA_LGG_65_MRI_RNAseq_tracing_evaluation_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"


def set_font(run, size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, 9.5, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_font(r, 9.5)
    set_table_widths(table, widths)
    doc.add_paragraph()


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_font(a, bold=True)
        b = p.add_run(text[len(bold_lead) :])
        set_font(b)
    else:
        r = p.add_run(text)
        set_font(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_font(r)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, bold=True, color=BLUE if level < 3 else DARK_BLUE)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, 9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after in ((1, 16, 16, 8), (2, 13, 12, 6), (3, 12, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if level < 3 else DARK_BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "TCGA-LGG MRI × RNA-seq 脑区溯源评估"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_font(run, 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("第 ")
    set_font(r, 9, color=MUTED)
    add_page_number(footer)
    r = footer.add_run(" 页")
    set_font(r, 9, color=MUTED)


def build() -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("技术评估报告")
    set_font(r, 10, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("TCGA-LGG 65例 MRI真值与RNA-seq脑区溯源评估")
    set_font(r, 23, bold=True, color=RGBColor(0, 0, 0))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("基于BraTS分割、SRI24/TZO116+图谱及现有主溯源路线")
    set_font(r, 13, color=MUTED)
    for label, value in (
        ("日期", "2026-06-09"),
        ("对象", "BraTS-TCGA-LGG 65名患者"),
        ("分割", "62例人工校正，3例自动分割"),
        ("输出", "Network、Lobe、Broad anatomy准确率与误差诊断"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        a = p.add_run(f"{label}：")
        set_font(a, 10.5, bold=True)
        b = p.add_run(value)
        set_font(b, 10.5)

    add_heading(doc, "执行结论")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(
        "当前路线不能可靠完成脑瘤所在Network的Top1定位。主分析Network Top1/Top3为13.8%/36.9%；"
        "Lobe和Broad anatomy的Top3超过80%，但主要来自候选范围扩展，不能等同于精确定位成功。"
    )
    set_font(r, 12, bold=True, color=DARK_BLUE)
    add_table(
        doc,
        ["层级", "Top1", "Top3", "Top1 95% CI", "Top3 95% CI"],
        [
            ["Network", "9/65，13.8%", "24/65，36.9%", "7.5%–24.3%", "26.2%–49.1%"],
            ["Lobe", "12/65，18.5%", "55/65，84.6%", "10.9%–29.6%", "73.9%–91.4%"],
            ["Broad anatomy", "10/65，15.4%", "54/65，83.1%", "8.6%–26.1%", "72.2%–90.3%"],
        ],
        [1800, 1500, 1500, 2280, 2280],
    )
    add_bullets(
        doc,
        [
            "主要失败模式是Subcortical预测塌缩：52/65患者RNA-seq Top1为Subcortical，MRI主导真值仅5/65。",
            "仅保留62例人工校正分割后，Network Top1/Top3为14.5%/38.7%，说明分割来源不是主要误差。",
            "该外部验证主要揭示肿瘤组织RNA相对健康脑图谱的疾病域偏移，尚不能证明RNA信号可准确定位真实肿瘤脑区。",
        ],
    )

    add_heading(doc, "1. 数据与真值生成")
    add_body(
        doc,
        "MRI来自BraTS-TCGA-LGG预处理四模态影像及GlistrBoost分割；RNA-seq使用本地已完成溯源的TCGA-LGG primary tumor样本；按TCGA patient barcode完成65例一一匹配。"
    )
    add_body(
        doc,
        "BraTS影像与SRI24图谱均为240×240×155、1 mm各向同性空间。根据NIfTI仿射矩阵将SRI24/TZO116+标签重采样到BraTS方向。肿瘤直接落入解剖标签的中位覆盖比例为66.7%，未直接标注的白质体素采用最近解剖标签传播。"
    )
    add_bullets(
        doc,
        [
            "Whole tumor：分割标签1、2、4并集，作为主分析。",
            "Tumor core：标签1和4；Edema：标签2；Tumor center：全肿瘤体素质心。",
            "严格规则：只接受重叠比例最高的单一标签。",
            "容错规则：接受占肿瘤体积至少20%的候选标签，处理跨脑叶或跨Network病灶。",
        ],
    )
    add_body(
        doc,
        "重要边界：本报告中的“真值”是专家/自动肿瘤分割结合标准图谱得到的MRI派生参考标签，并非放射科医师逐例手工确认的临床金标准。",
        bold_lead="重要边界：",
    )

    add_heading(doc, "2. 主结果")
    add_table(
        doc,
        ["匹配规则", "Network Top1", "Network Top3", "Lobe Top1", "Lobe Top3", "Broad Top1", "Broad Top3"],
        [
            ["严格单标签", "7.7%", "23.1%", "10.8%", "80.0%", "9.2%", "73.8%"],
            ["容错多标签", "13.8%", "36.9%", "18.5%", "84.6%", "15.4%", "83.1%"],
        ],
        [1800, 1260, 1260, 1260, 1260, 1260, 1260],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUTDIR / "brats_tcga_lgg_65_primary_accuracy.png"), width=Inches(6.25))
    c = doc.add_paragraph("图1. 全肿瘤容错规则下的主准确率及Wilson 95%置信区间")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        set_font(run, 9, color=MUTED)

    add_heading(doc, "3. 敏感性分析")
    add_table(
        doc,
        ["真值口径", "Network Top1", "Network Top3", "Lobe Top1", "Lobe Top3"],
        [
            ["Whole tumor，容错", "13.8%", "36.9%", "18.5%", "84.6%"],
            ["Tumor core，容错", "16.9%", "46.2%", "20.0%", "84.6%"],
            ["Edema，容错，n=64", "25.0%", "45.3%", "31.3%", "92.2%"],
            ["Tumor center，单标签", "21.5%", "30.8%", "23.1%", "52.3%"],
            ["Whole tumor，人工分割n=62", "14.5%", "38.7%", "19.4%", "85.5%"],
        ],
        [3000, 1590, 1590, 1590, 1590],
    )
    add_body(
        doc,
        "不同真值口径下Network准确率均偏低，结论稳定。Tumor core的Network Top3略高于Whole tumor，但仍不足以支持可靠Network定位。"
    )

    add_heading(doc, "4. 误差模式")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUTDIR / "brats_tcga_lgg_65_network_distribution.png"), width=Inches(6.15))
    c = doc.add_paragraph("图2. MRI主导Network与RNA-seq Top1预测分布")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        set_font(run, 9, color=MUTED)
    add_bullets(
        doc,
        [
            "52/65（80.0%）患者Top1为Subcortical。",
            "最常见Top3组合为“Subcortical + Hippocampal formation + OMPFC”，出现45/65（69.2%）。",
            "MRI真值最多的Network为Lateral PFC，18/65（27.7%）；当前严格Network Top1为7.7%，低于简单多数类基线。",
            "Top1/Top2 margin最高四分位的17例中，Network Top1仍仅23.5%，现有margin筛选能力有限。",
        ],
    )

    add_heading(doc, "5. 结果解释")
    add_body(
        doc,
        "TCGA RNA-seq来自肿瘤组织，而非正常脑组织或血浆cfRNA。表达矩阵包含大量与位置无关或弱相关的增殖、缺氧、免疫浸润、坏死、细胞外基质和反应性胶质程序。直接投射到健康猕猴脑图谱时，跨物种差异与疾病域偏移会共同主导得分。"
    )
    add_bullets(
        doc,
        [
            "健康参考交叉验证表现较好，不代表可直接迁移到脑瘤组织RNA。",
            "本次外部验证明确暴露了Subcortical先验偏置。",
            "Lobe/Broad Top3可作为候选覆盖率，但不应表述为精确脑区定位准确率。",
        ],
    )

    add_heading(doc, "6. 下一轮优化路线")
    add_table(
        doc,
        ["优先级", "改进项", "实施要点"],
        [
            ["P0", "肿瘤域去偏", "对细胞周期、免疫、基质、缺氧、血管生成和泛胶质反应基因降权或回归。"],
            ["P0", "层级分类", "先预测lobe/broad anatomy，再在候选脑叶内预测Network，限制全局Subcortical压制。"],
            ["P0", "类别先验校准", "在65例上使用嵌套LOPO/分层交叉验证估计偏置项，患者严格隔离。"],
            ["P1", "保守区域标记集", "筛选人-猕猴一对一同源且在两物种健康脑中均区域特异的基因。"],
            ["P1", "Core/edema双通路", "分别建模肿瘤核心与周边反应，评估两类信号的空间贡献。"],
            ["P1", "置信度重建", "联合熵、校准后margin、重采样稳定性与基因子集一致性。"],
        ],
        [900, 2100, 6360],
    )
    add_body(
        doc,
        "下一轮主指标建议采用Network macro-recall、balanced accuracy、Top1/Top3和校准误差；Lobe/Broad Top3仅作为覆盖率辅助指标。"
    )

    add_heading(doc, "7. 输出文件")
    add_bullets(
        doc,
        [
            "brats_tcga_lgg_65_mri_truth_and_predictions.csv：65例逐患者真值、预测、margin和命中状态。",
            "brats_tcga_lgg_65_accuracy_metrics.csv：全部真值口径的准确率和Wilson 95%置信区间。",
            "network_top1_confusion_matrix.csv：Network Top1混淆矩阵。",
            "sri24_tzo116_resampled_to_brats.nii.gz：BraTS空间图谱标签。",
            "scripts/evaluate_brats_tcga_lgg_65_mri_truth.py：可复现评估脚本。",
        ],
    )

    add_heading(doc, "参考资料")
    add_body(
        doc,
        "1. Rohlfing T, et al. The SRI24 multichannel atlas of normal adult human brain structure. Human Brain Mapping. 2010;31:798–819. https://pmc.ncbi.nlm.nih.gov/articles/PMC2915788/"
    )
    add_body(
        doc,
        "2. Bakker R, et al. The Scalable Brain Atlas: instant web-based access to public brain atlases and related content. Neuroinformatics. 2015. https://pmc.ncbi.nlm.nih.gov/articles/PMC4469098/"
    )
    add_body(doc, "3. BraTS Challenge, CBICA, University of Pennsylvania. https://www.med.upenn.edu/cbica/brats/")

    doc.core_properties.title = "TCGA-LGG 65例 MRI真值与RNA-seq脑区溯源评估"
    doc.core_properties.subject = "BraTS MRI-derived truth and RNA-seq tracing evaluation"
    doc.core_properties.author = "Brain tracing project"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
