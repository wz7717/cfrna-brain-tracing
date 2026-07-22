"""Apply the in-scope TCGA/BraTS denominator correction with local DOCX edits."""

from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MAIN = ROOT / "0711Bioinformatics_Application_Note_cfRNA_BrainTrace_submission_P0_review_metrics_figure1_cluster_bh_clarified.docx"
SUPP = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_complete_with_sparse_donor_tables_cluster_bh_clarified.docx"


def replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        raise ValueError(f"Expected text missing: {old[:90]!r}")
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    raise ValueError("Expected replacement crosses Word runs; refusing a format-damaging edit")


def replace_in_cell(cell, old: str, new: str) -> None:
    for paragraph in cell.paragraphs:
        if old in paragraph.text:
            replace_in_paragraph(paragraph, old, new)
            return
    raise ValueError(f"Expected cell text missing: {old!r}")


def main() -> None:
    main_doc = Document(MAIN)
    p = main_doc.paragraphs[26]
    replace_in_paragraph(
        p,
        "In that setting, Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%.",
        "After exclusion of one MRI-derived cerebellar out-of-scope case, Network Top3 was 40.63% (26/64) and broad-anatomy Top3 was 65.63% (42/64).",
    )
    replace_in_paragraph(
        main_doc.paragraphs[48],
        "TCGA/BraTS Network and broad-anatomy Top3 values are 40.00% and 64.62%.",
        "TCGA/BraTS Network and broad-anatomy Top3 values are 40.63% and 65.63% among 64 in-scope cases.",
    )
    main_doc.save(MAIN)

    supp_doc = Document(SUPP)
    p = supp_doc.paragraphs[39]
    old = (
        "In TCGA/BraTS, hybrid Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%, supporting only coarse anatomical consistency. "
        "Against a nominal 30% Network Top3 reference level, the TCGA/BraTS Network Top3 value of 26/65 has a one-sided normal-approximation p value of approximately 0.039; "
        "the more conservative exact binomial test gives p=0.0548. Against the endpoint-specific weighted random baseline of 20.28%, the exact one-sided binomial p value is 2.17e-4."
    )
    new = (
        "In TCGA/BraTS, one MRI-derived cerebellar out-of-scope case was excluded, leaving 64 evaluable patients. "
        "Hybrid Network Top3 was 40.63% (26/64) and broad-anatomy Top3 was 65.63% (42/64), supporting only coarse anatomical consistency. "
        "Against a nominal 30% Network Top3 reference level, the value of 26/64 has a one-sided normal-approximation p value of approximately 0.032; "
        "the more conservative exact binomial test gives p=0.0455. Against the endpoint-specific weighted random baseline of 20.28%, the exact one-sided binomial p value is 1.61e-4."
    )
    replace_in_paragraph(p, old, new)
    replace_in_cell(supp_doc.tables[3].rows[2].cells[2], "65 patients", "65 total; 64 in-scope evaluable")
    replace_in_cell(supp_doc.tables[4].rows[6].cells[3], "15.38%", "15.63%")
    replace_in_cell(supp_doc.tables[4].rows[6].cells[4], "40.00%", "40.63%")
    replace_in_cell(supp_doc.tables[4].rows[7].cells[3], "13.85%", "14.06%")
    replace_in_cell(supp_doc.tables[4].rows[7].cells[4], "64.62%", "65.63%")
    supp_doc.save(SUPP)
    print("TCGA/BraTS in-scope denominator correction applied")


if __name__ == "__main__":
    main()
