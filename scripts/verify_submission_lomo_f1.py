#!/usr/bin/env python
"""Verify that Supplement Table S12 and key prose use formal LOMO F1 values."""

from __future__ import annotations

import argparse
import csv
import json
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.lomo_f1 import CANONICAL_FORMAL_PATH, compute_lomo_network_metrics, load_formal_predictions  # noqa: E402


def load_detail(path: Path) -> dict[str, dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return {row["class"]: row for row in csv.DictReader(handle)}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--supplement", type=Path, required=True)
    parser.add_argument("--detail", type=Path, default=ROOT / "reproducibility" / "formal_lomo_network_f1.csv")
    args = parser.parse_args()

    detail = load_detail(args.detail)
    expected = compute_lomo_network_metrics(load_formal_predictions(CANONICAL_FORMAL_PATH))
    if len(detail) != 10:
        raise SystemExit(f"Expected 10 formal class rows, got {len(detail)}")

    document = Document(args.supplement)
    table = next(
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text == "Network"
        and table.rows[0].cells[-1].text == "LOMO F1"
    )
    checked = 0
    for row in table.rows[1:]:
        cls = row.cells[0].text
        if cls not in detail:
            raise SystemExit(f"Table S12 class not found in formal detail: {cls}")
        source = detail[cls]
        expected_cells = [
            f"{float(source['precision']):.2f}",
            f"{float(source['recall']):.2f}",
            f"{float(source['f1']):.2f}",
        ]
        actual_cells = [cell.text for cell in row.cells[5:8]]
        if actual_cells != expected_cells:
            raise SystemExit(f"Table S12 mismatch for {cls}: {actual_cells} != {expected_cells}")
        checked += 1
    if checked != 10:
        raise SystemExit(f"Expected 10 Table S12 rows, got {checked}")

    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = [
        "in the formal LOMO prediction set, 41/41 predicted Subcortical calls were correct",
        "anchored to the LOSO Subcortical result",
        "Operculum/Insula had precision 0.43, recall 0.57 and F1 0.49 in LOSO, versus precision 0.35, recall 0.60 and F1 0.45",
        "not presented as absolute extrema across Networks",
        "formal_lomo_network_f1.csv",
    ]
    missing = [text for text in required if text not in paragraphs]
    if missing:
        raise SystemExit(f"Required corrected Supplement text is missing: {missing}")
    forbidden = [
        "LOMO (precision 0.41, recall 0.68, F1 0.51)",
        "both LOSO and LOMO (42/42 predictions correct)",
        "Parietal Network exhibits the lowest precision",
    ]
    present = [text for text in forbidden if text in paragraphs]
    if present:
        raise SystemExit(f"Superseded Supplement text remains: {present}")

    print(json.dumps({"table_s12_rows_checked": checked, "formal_micro_f1": expected["summary"]["micro_f1"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
