#!/usr/bin/env python3
"""Apply non-Huang provenance corrections to submission DOCX copies.

The script accepts the current submission documents as input and writes new
copies. It replaces only identified non-Huang scientific-provenance text,
preserving the first run's character formatting and the original paragraph/table
structure.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def replace_once(document: Document, marker: str, replacement: str, label: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if marker in paragraph.text]
    if len(matches) != 1:
        raise ValueError(f"{label}: expected one paragraph containing {marker!r}, found {len(matches)}")
    paragraph = matches[0]
    updated = paragraph.text.replace(marker, replacement)
    if updated == paragraph.text:
        raise ValueError(f"{label}: replacement made no change")
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def replace_table_cell_once(document: Document, marker: str, replacement: str, label: str) -> None:
    matches = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if marker in cell.text:
                    matches.append(cell)
    if len(matches) != 1:
        raise ValueError(f"{label}: expected one table cell containing {marker!r}, found {len(matches)}")
    cell = matches[0]
    paragraph = cell.paragraphs[0]
    updated = cell.text.replace(marker, replacement)
    if len(cell.paragraphs) != 1:
        raise ValueError(f"{label}: expected one paragraph in target cell")
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def update_main(document: Document) -> None:
    replace_once(
        document,
        "Ensembl orthology mapped 188/200 Network genes and 5,324/8,800 region-signature rows.",
        "Ensembl orthology mapped 188/200 Network genes and 5,324/8,800 gene-by-region row occurrences; 3,476/8,800 row occurrences lacked a direct high-confidence human ortholog under the frozen mapping rule. These are not independent gene counts.",
        "main orthology unit",
    )
    replace_once(
        document,
        "AHBA human RNA-seq used two donors and yielded 223 Network- and 88 group/exact-evaluable samples;",
        "AHBA human RNA-seq used two donors and had 231 replicate-collapsed tissues; endpoint-specific evaluability was 223 Network and 88 group/exact samples, not a unique sequential attrition pipeline;",
        "main AHBA endpoint wording",
    )
    replace_once(
        document,
        "candidate-set any-hit Network Top3 36.51% (23/63; descriptive sensitivity).",
        "candidate-set any-hit Network Top3 36.51% (23/63; descriptive sensitivity). Across strict Top3 truth bases, Network accuracy ranged from 15.38% to 29.23% (13.85 percentage points), and broad-anatomy accuracy ranged from 49.23% to 82.54% (33.31 percentage points).",
        "main TCGA truth-basis range",
    )


def update_supplement(document: Document) -> None:
    replace_once(
        document,
        "In the frozen 8,800-row region-signature audit, 5,324 rows had a humanized symbol and 3,476 ENSMFAG rows remained unmapped.",
        "In the frozen 8,800-row region-signature audit, 5,324/8,800 gene-by-region row occurrences had a humanized symbol and 3,476/8,800 row occurrences remained unmapped under the frozen mapping rule.",
        "supplement orthology row unit",
    )
    replace_once(
        document,
        "The 60.5% row-level humanization rate means that approximately 39.5% of macaque reference genes lack a direct human ortholog in the frozen lookup table.",
        "The 60.5% row-level humanization rate means that 3,476/8,800 gene-by-region row occurrences (39.5%) lacked a direct high-confidence human ortholog under the frozen mapping rule; this is not a percentage of independent macaque reference genes.",
        "supplement orthology denominator",
    )
    replace_once(
        document,
        "Technical replicates were collapsed by donor and tissue-sample identifier through raw-count summation before logCPM (231 independent tissues);",
        "Technical replicates were collapsed by donor and tissue-sample identifier through raw-count summation before logCPM (231 replicate-collapsed tissues); endpoint-specific evaluability was 223 Network and 88 group/exact samples, not a unique sequential attrition pipeline;",
        "supplement AHBA method wording",
    )
    replace_once(
        document,
        "AHBA sample cascade. The AHBA validation pipeline reduced 231 independent tissue samples (after collapsing technical RNA-seq replicates by donor and tissue-sample identifier through raw-count summation before logCPM) to 223 Network-evaluable samples and 88 resolution-group/exact-region-evaluable samples through two sequential filters. Step 1 (231 to 223): eight samples were excluded because their AHBA anatomical labels could not be harmonized to any of the 10 macaque-derived Network labels. This filter removes samples whose primary tissue identity falls outside the coarse cortical/subcortical scope of the Bo2023-derived hierarchy (e.g., white matter, ventricular zone, or structures without a supported cross-species label mapping). The 3.5% exclusion rate is consistent with a conservative label-harmonization policy that retains samples only when a defensible macaque-to-human label bridge exists in the frozen Saleem crosswalk.",
        "AHBA endpoint evaluability. The AHBA validation source contains 231 replicate-collapsed tissue samples. Network evaluability is 223 because eight samples have no valid Network mapping; resolution-group and exact-region evaluability are each 88 because those endpoints require a supported exact-region mapping. These are endpoint-specific eligibility counts, not two sequential filters in one unique scientific sample-flow pipeline. The eight non-Network-evaluable samples lie outside the supported coarse cortical/subcortical label bridge (for example white matter, ventricular zone, or structures without a defensible frozen crosswalk mapping).",
        "supplement AHBA cascade replacement",
    )
    replace_once(
        document,
        "Step 2 (223 to 88): 135 of the 223 Network-mapped samples lacked a supported resolution-group or exact-region label in the macaque-derived hierarchy. This is not a sample-quality filter but a label-support filter: the AHBA human brain atlas uses a different anatomical nomenclature and sampling density from the Bo2023 macaque atlas, so many human tissue locations cannot be mapped to a specific resolution group or exact region after the frozen crosswalk. The resulting 88-sample subset (39.5% retention from 223; 38.1% from 231) supports all three hierarchical levels and is used for resolution-group and exact-region validation. Network-level results are reported on the full 223-sample set. The single-label sensitivity analyses (n=56 for Network, n=40 for group/exact) are a further within-subset restriction and are not additional pipeline steps. All steps use the frozen label crosswalk only; no sample was reclassified post hoc. Complete per-sample trace data are archived in v4_p0_5_ahba_trace.csv.",
        "The 135 Network-evaluable samples that are not group/exact-evaluable lack a supported exact-region mapping in the macaque-derived hierarchy. This is a label-support limitation, not a sample-quality filter: AHBA uses different nomenclature and sampling density, so many human tissue locations cannot receive a specific frozen resolution-group or exact-region mapping. The 88 exact-mapped samples are used for the resolution-group and exact-region endpoints, whereas Network results use all 223 Network-evaluable samples. The single-label sensitivities (Network n=56; group/exact n=40) are within-endpoint restrictions, not additional sample-flow stages. The canonical per-sample source is `ahba_endpoint_evaluability_ledger.csv`; the retained v4 AHBA traces are historical engineering traces and not canonical endpoint accounting.",
        "supplement AHBA nonsequential clarification",
    )
    replace_once(
        document,
        "while broad Top3 was 49.23%, 69.23%, 82.54% and 70.77% (range 32.02 points).",
        "while broad Top3 was 49.23%, 69.23%, 82.54% and 70.77% (range 33.31 percentage points).",
        "supplement TCGA broad range",
    )
    replace_once(
        document,
        "Conditional on a correct Network set, LOSO exact-region Top3 was 49.07% and resolution-group Top3 was 78.32%; after a missed Network set, LOSO exact-region recovery was 0.00%.",
        "Within the same 814 exact-evaluable samples, the Network candidate set retained truth in 750 and missed it in 64; conditional exact-region/resolution-group Top3 was 368/750=49.07% and 590/750=78.67%, respectively, and recovery after a Network candidate-set miss was 0%.",
        "supplement tier conditional metric",
    )
    replace_once(
        document,
        "In the frozen cascade, approximately 66 of 446 exact misses (14.8%) coincided with fixed Network-candidate-set misses; exact and Group Top3 conditional on a correct candidate set were 49.07% and 78.32%, and recovery after a Network-set miss was 0%.",
        "In the frozen cascade, 64 of 446 exact Top3 misses (14.35%) coincided with Network candidate-set misses within the same 814 exact-evaluable samples; conditional Exact/Group Top3 was 368/750=49.07% and 590/750=78.67%, and recovery after a Network candidate-set miss was 0%.",
        "supplement tier cascade denominator",
    )
    replace_table_cell_once(
        document,
        "231 independent RNA-seq tissue samples after technical-replicate collapse (223 Network-qualified; 88 group/exact evaluable)",
        "231 replicate-collapsed RNA-seq tissue samples; endpoint-specific evaluability: Network n=223 and group/exact n=88",
        "supplement AHBA table sample count",
    )
    replace_table_cell_once(
        document,
        "231 independent; Network n=223; group/exact n=88",
        "231 replicate-collapsed tissues; Network n=223; group/exact n=88 (endpoint-specific)",
        "supplement AHBA table endpoint wording",
    )


def render_copy(input_path: Path, output_path: Path, updater) -> None:
    document = Document(input_path)
    updater(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--main-input", type=Path, required=True)
    parser.add_argument("--supplement-input", type=Path, required=True)
    parser.add_argument("--main-output", type=Path, required=True)
    parser.add_argument("--supplement-output", type=Path, required=True)
    args = parser.parse_args()
    render_copy(args.main_input, args.main_output, update_main)
    render_copy(args.supplement_input, args.supplement_output, update_supplement)
    print(args.main_output)
    print(args.supplement_output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
