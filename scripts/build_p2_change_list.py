#!/usr/bin/env python3
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

out=Path("manuscript/BrainTrace_v5_round4_P2_修改清单.docx")
repo_root = Path(__file__).resolve().parents[1]
tier = json.loads((repo_root / "reproducibility" / "tier_cascade_loso_summary.json").read_text(encoding="utf-8"))
beam = tier["network_candidate_miss_share_of_exact_top3_misses"]
exact_conditional = tier["exact_top3_given_network_truth_retained"]
group_conditional = tier["group_top3_given_network_truth_retained"]
d=Document(); s=d.sections[0]
s.top_margin=s.bottom_margin=Inches(.65); s.left_margin=s.right_margin=Inches(.7)
d.styles["Normal"].font.name="Arial"; d.styles["Normal"].font.size=Pt(9)
d.styles["Title"].font.name="Arial"; d.styles["Title"].font.size=Pt(18)
p=d.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("BrainTrace v5 第四轮：P2 类问题修改清单")
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("范围：报告称28项，但仅明确列出15项；本轮不推断未列出的13项").italic=True

d.add_heading("一、核心计算结果",level=1)
for x in [
"冻结输出层消融：Full Network/Group/Exact Top3=91.94%/72.48%/45.21%；移除某一显示层不重训其余模型，故这是信息移除诊断而非新模型。",
 f"Network candidate-set miss约占Exact错误{beam['correct']}/{beam['n']}={beam['percent']:.2f}%；"
 f"candidate set包含truth时Exact/Group Top3={exact_conditional['percent']:.2f}%/"
 f"{group_conditional['percent']:.2f}%，candidate-set miss后恢复率0%。",
"匹配RF（同一冻结Network truth、固定200基因、300 trees、balanced_subsample、min leaf=2、9折LOMO完整重训）：Top1 389/819=47.50%，Top3 680/819=83.03%。",
"Exact F1：LOSO median=0.1538，IQR=0.2857；LOMO median=0.13245，IQR=0.17895。",
"探索性富集：446个GO-BP、11个KEGG FDR显著条目；细胞类型仅ExcN（q=2.91e-19）及源标签IhnN（q=2.74e-5）通过BH 0.05。",
]:
 d.add_paragraph(x,style="List Bullet")

items=[
("1","三层架构消融","完成冻结模型输出层诊断和beam错误归因；明确非重训消融。","主文P2段；补充S20"),
("2","主文缺工具架构图","确认Figure 1已为完整工具架构/流程/验证图，并更新标题措辞。","Figure 1及Supplementary Figure S1"),
("3","GO/KEGG和细胞类型偏倚","按200 panel、21,668背景开展探索性FDR分析；强调事后性和非机制结论。","主文P2段；补充S20；CSV"),
("4","Bo2023-Saleem对应表","生成110行region/Network/Saleem式全名/lobe/broad-map交叉表。","P2_Bo2023_Saleem_crosswalk.csv"),
("5","Claustrum/Piriform分类论证","明确二者是锁定转录组/atlas操作性成员，不宣称经典皮层下核团。","主文及S20"),
("6","BioMart版本/日期","冻结查询信息未归档，不猜测；记录2026-07-30重查时Ensembl 116，未替换冻结映射。","主文及S20"),
("7","5-NN余弦距离理由","说明其强调表达方向、弱化全局幅度；未进行距离选择搜索。","主文及S20"),
("8","Live demo限制","确认无应用侧rate limiter或上传覆盖；host限制适用；报告153×28,415和222 MiB基准。","主文及S20"),
("9","三个requirements命名","明确minimal app/full reproduction/testing extras三种角色，不宣称可互换。","主文及S20"),
("10","F1 median/IQR","补报Exact和Network的LOSO/LOMO中位数与IQR。","主文及S20"),
("11","RF同200 panel","保留历史k=500并新增匹配200-gene RF；不替换正式路线。","主文、S20、Figure S2"),
("12","Friedman n=9近似","并列χ²近似与19,683模式精确枚举，限制结论强度。","主文及S20"),
("13","Sign-flip exchangeability","补充donor间独立、零假设下符号对称；非sample-row置换。","主文及S20"),
("14","Cerebellum排除理由","明确训练参考、层级和验证truth均缺失，故必须abstain/out-of-scope。","主文及S20"),
("15","Donor confusion可视化","新增匹配200-gene RF的9 donor小多图；明确并非正式路线。","Supplementary Figure S2"),
]
d.add_heading("二、逐项审阅清单",level=1)
t=d.add_table(rows=1,cols=4); t.style="Table Grid"
for c,h in zip(t.rows[0].cells,["项","问题","处理","位置"]): c.text=h
for row in items:
 for c,v in zip(t.add_row().cells,row): c.text=v

d.add_heading("三、主要边界与可复现产物",level=1)
for x in [
"所有P2新分析均为post hoc descriptive/exploratory，不改变冻结正式路线。",
"GO/KEGG结果高度受嵌套ontology、数据库版本和背景集影响，不作为细胞来源或机制证据。",
"Figure S2展示匹配RF而非正式层级路线，避免将新敏感性图误归因于主模型。",
"计算目录：manuscript/calculations/p2；总索引：P2_audit.json。",
"结构QA：trackRevisions开启、无嵌套修订；版式QA覆盖主文、补充材料和本清单全部页面。"
]: d.add_paragraph(x,style="List Bullet")

d.add_heading("四、建议审阅重点",level=1)
d.add_paragraph("重点确认：(1) 接受冻结输出层诊断而非重新训练消融；(2) 匹配RF作为敏感性而不替代历史结果；"
                "(3) 457个富集条目仅作探索性注释；(4) BioMart冻结版本缺失被透明披露；"
                "(5) Figure S2的模型身份和解释边界足够醒目。")
d.save(out); print(out)
