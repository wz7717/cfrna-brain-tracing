"""Update only supplementary-methods/results prose for the P0 reviewer analyses."""

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document


ROOT = Path(r"D:\Download\cfrna-brain-tracing-streamlit-cloud-ready")
SOURCE = ROOT / "Bioinformatics_Application_Note_Supplementary_File_v6_references_checked.docx"
OUTPUT = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_methods_results_updated.docx"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace prose while retaining paragraph formatting and the first run style."""
    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for run in paragraph.runs:
        paragraph._p.remove(run._r)
    run = paragraph.add_run(new_text)
    if first_rpr is not None:
        run._r.insert(0, first_rpr)


def main() -> None:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    doc = Document(source)

    replacements = {
        8: (
            "Internal validation used the Bo2023 macaque brain reference. Two settings were evaluated. "
            "Leave-one-sample-out validation tested whether the route could recover the held-out sample label when other samples from the reference remained available. "
            "Leave-one-monkey-out validation held out all samples from one animal at a time and therefore tested donor-level generalization. "
            "Network metrics included all 819 samples because every held-out sample retained a supported Network label. Resolution-group and exact-region metrics required the truth region to remain represented in the corresponding training fold. "
            "Five LOSO samples and seven LOMO samples did not meet that region-reference requirement and were excluded only from region-level denominators. "
            "Top1, Top3 and median true-rank were used to distinguish exact calls from candidate-list recovery. "
            "To account for repeated samples from the same animals, donor-macro accuracy, donor-level bootstrap intervals and small-cluster-corrected sandwich intervals used monkey as the clustering unit. "
            "Sensitivity to tissue-to-cfRNA sparsity was assessed without retuning the locked route by strict LOSO query perturbation: only the held-out count vector was abundance-sampled and then binomially thinned, while the training reference, marker panels, hierarchy and route parameters remained fixed. "
            "Mild, moderate, severe and extreme scenarios targeted 80%, 60%, 40% and 20% detected-gene retention at 50%, 20%, 5% and 1% read-depth fractions, respectively; each non-baseline scenario used three random replicates and donor-level bootstrap across the nine macaques."
        ),
        12: (
            "External validation was limited by the label resolution available in each dataset. AHBA human brain RNA-seq was used for mapped-label validation because its anatomical labels could be harmonized to Network, resolution-group and a subset of exact-region labels. "
            "Technical AHBA RNA-seq replicates were collapsed by donor and tissue-sample identifier through raw-count summation before logCPM calculation, yielding 231 independent tissue samples. "
            "TCGA/BraTS glioma tissue RNA-seq with MRI-derived labels was used only for coarse anatomical consistency because its truth labels are human imaging labels, not Bo2023 macaque exact-region identifiers. "
            "GSE189919 was used to test whether an external matrix could be projected into the model gene space; it was not used for accuracy estimation because patient-level anatomical truth was unavailable."
        ),
        14: (
            "Label harmonization was performed only to the level supported by each dataset. AHBA anatomical labels were mapped to anatomically harmonized allowed-label sets in the macaque-derived Network, resolution-group and exact-region hierarchy where a supported mapping existed. "
            "Fine regional references were Network-qualified using network_id::region_id, preventing duplicate Bo2023 region identifiers from being merged across Networks. "
            "Results are interpreted as mapped-label transfer rather than direct anatomical equivalence; mapping multiplicity is reported because an AHBA sample can support more than one allowed Network or Network-qualified region. "
            "TCGA/BraTS labels were derived from human MRI/tumour context and therefore support only coarse tumour-tissue anatomical consistency. Biofluid datasets without patient-level anatomical truth were not used for localization accuracy and are reported only as projection-feasibility or transfer stress tests."
        ),
        31: (
            "The reviewer-response evidence package is organized into three tiers. P0 hard-evidence files contain the formal validation metrics, Wilson confidence intervals, random baselines, binomial tests, donor-macro and donor-level bootstrap analyses, cluster-robust LOSO-versus-LOMO comparisons with Benjamini-Hochberg correction, sparse-query sensitivity analyses, confusion matrices, class-level F1, denominator audits, marker-methodology audits, Network anatomy tables, resolution-group hierarchy tables and AHBA technical-replicate and mapping-granularity summaries. "
            "P1 diagnostics contain dual-space consistency, error-cascade and confidence-diagnostic analyses. P2 completeness files contain development/comparator materials, including same-field tool positioning, simple ML baselines, test coverage, locked dependency files and random-seed registry. "
            "P2 comparator files are retained as development, comparator, sensitivity or diagnostic evidence; they are not merged into the formal three-tier hybrid validation endpoint."
        ),
        36: (
            "In the complete LOSO validation, the formal route achieved Network Top1/Top3 of 58.24%/92.19% across all 819 samples. Resolution-group and exact-region Top3 were 72.36% and 45.33% among 814 reference-supported samples. "
            "In complete LOMO validation, Network Top3 was 91.21% across all 819 samples, while resolution-group and exact-region Top3 were 69.09% and 42.36% among 812 reference-supported samples. "
            "The 92.19% LOSO Network Top3 value uses all 819 Network-evaluable samples as the denominator. Region-level LOSO metrics use 814 reference-supported samples because five samples lacked a truth-region reference after fold construction. Region-level LOMO metrics use 812 reference-supported samples because seven samples lacked a truth-region reference after fold construction. "
            "The earlier LOSO Network value of 92.38% was conditional on the 814 region-evaluable samples and is retained only as a legacy denominator inconsistency, not as the submission result. "
            "On common evaluable samples, donor-macro Top3 accuracy was 89.69% versus 88.51% for Network, 68.63% versus 67.14% for resolution group, and 42.19% versus 39.11% for exact region in LOSO and LOMO, respectively. "
            "Median true-rank increased from Network to exact-region levels, consistent with decreasing anatomical certainty at finer resolution. Resolution group is therefore the preferred region-level endpoint, while exact-region output is retained as a candidate ranking."
        ),
        37: (
            "Wilson confidence intervals (Wilson, 1927) and statistical tests were generated for the formal P0 evidence package. Network Top3 was 92.19% in LOSO (755/819; 95% CI 90.14-93.83%) and 91.21% in LOMO (747/819; 95% CI 89.07-92.96%). "
            "Resolution-group Top3 was 72.36% in LOSO (589/814; 95% CI 69.19-75.32%) and 69.09% in LOMO (561/812; 95% CI 65.83-72.17%). Exact-region Top3 was 45.33% in LOSO (369/814; 95% CI 41.94-48.77%) and 42.36% in LOMO (344/812; 95% CI 39.01-45.79%). "
            "Each endpoint exceeded its uniform random baseline by one-sided binomial testing, including exact-region LOSO Top3 versus 7.92% uniform random expectation (p=4.43e-181). "
            "The small-cluster-corrected LOSO-minus-LOMO differences were 0.98 percentage points for Network (95% CI 0.17-1.78; unadjusted p=0.0230), 3.33 percentage points for resolution group (95% CI -0.09 to 6.74; p=0.0552), and 2.96 percentage points for exact region (95% CI -1.87 to 7.79; p=0.1959). "
            "After Benjamini-Hochberg correction across the three prespecified hierarchical Top3 comparisons, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold. "
            "The original paired exact McNemar p values were 0.1686, 0.0389 and 0.0532 for Network, resolution group and exact region, respectively; the corresponding BH-adjusted values were 0.1686, 0.0798 and 0.0798."
        ),
        39: (
            "In AHBA mapped-label external validation, technical RNA-seq replicates were collapsed by donor and tissue-sample identifier through raw-count summation before logCPM calculation, yielding 231 independent tissue samples. "
            "Fine regional references were Network-qualified (network_id::region_id). The formal three-tier hybrid route achieved Network Top1/Top3 of 73.99% (165/223; 95% CI 67.86-79.31%)/94.62% (211/223; 90.83-96.90%), resolution-group Top1/Top3 of 40.91% (36/88; 31.23-51.35%)/62.50% (55/88; 52.06-71.89%), and exact-region Top1/Top3 of 23.86% (21/88; 16.17-33.74%)/46.59% (41/88; 36.53-56.94%). "
            "Exact Top3 remained above both logCPM baseline and projected-VSD-only scoring. These are mapped-label transfer metrics, not unique anatomical-truth or MRI-derived localization accuracy. "
            "Mapping multiplicity was substantial: 74.9% of supported samples allowed more than one Bo2023 Network, and 54.5% of exact-mapped samples allowed more than one Network-qualified region. In the unique-Network subset, Network Top3 accuracy was 78.57% (44/56; 95% CI 66.18-87.29%). "
            "Thus, AHBA supports cross-species mapped-label transfer at coarse resolution but does not establish unique human-to-macaque exact-region equivalence. "
            "In TCGA/BraTS, hybrid Network Top3 was 40.00% and broad-anatomy Top3 was 64.62%, supporting only coarse anatomical consistency. Against a nominal 30% Network Top3 reference level, the TCGA/BraTS Network Top3 value of 26/65 has a one-sided normal-approximation p value of approximately 0.039; the more conservative exact binomial test gives p=0.0548. "
            "Against the endpoint-specific weighted random baseline of 20.28%, the exact one-sided binomial p value is 2.17e-4. These values support only coarse tumour-tissue consistency and do not convert TCGA/BraTS into an exact localization validation set. "
            "GSE189919 overlapped 15,622/21,668 projector genes, corresponding to 72.10% gene-space coverage, supporting projection feasibility rather than source-localization accuracy."
        ),
        43: (
            "The three prespecified internal hierarchical Top3 comparisons were corrected for multiple testing using the Benjamini-Hochberg procedure. "
            "The adjusted p values for the paired exact McNemar comparisons were 0.1686 for Network and 0.0798 for both resolution group and exact region; thus, no LOSO-versus-LOMO difference reached the conventional 0.05 threshold after correction. "
            "Endpoint-specific random-baseline tests and coarse external-consistency analyses address separate descriptive questions and are reported with their stated interpretation boundaries."
        ),
        51: (
            "Several layers of transfer separate the training reference from the intended cfRNA use case. First, the reference is tissue RNA-seq, while cfRNA reflects extracellular RNA abundance after cell release, degradation and clearance. Second, cfRNA contains contributions from multiple tissues and blood components, so the brain signal may be diluted. Third, external cfRNA studies often lack patient-level anatomical truth, making apparent localization impossible to validate directly. "
            "In strict locked-route LOSO sparse-query simulations, actual detected-gene coverage declined from 99.96% at baseline to 79.17%, 58.97%, 38.58% and 18.27% across mild-to-extreme scenarios. Network Top3 was relatively preserved under mild sparsity (91.25% versus 92.19% at baseline), but declined to 82.62%, 71.10% and 59.46% under moderate, severe and extreme sparsity. "
            "Resolution-group and exact-region Top3 declined more rapidly, from 72.36% and 45.33% at baseline to 40.34% and 17.49% under extreme sparsity. These simulations define a tissue-derived sparse-input sensitivity boundary, not clinical cfRNA localization accuracy. "
            "We therefore separate evidence types: internal Bo2023 LOSO/LOMO measures reference-domain traceability; AHBA measures mapped-label transfer to normal human brain RNA-seq; TCGA/BraTS measures coarse consistency in tumour tissue with MRI-derived labels; and biofluid datasets measure gene coverage and projection feasibility only. "
            "The required future validation is a prospective or curated cfRNA cohort with matched imaging or surgical anatomical truth, reported at the same Network/resolution-group/exact-region levels and with abstention for unsupported labels."
        ),
    }

    for index, text in replacements.items():
        replace_paragraph_text(doc.paragraphs[index], text)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Updated {len(replacements)} supplementary prose paragraphs")
    print(output)


if __name__ == "__main__":
    main()
