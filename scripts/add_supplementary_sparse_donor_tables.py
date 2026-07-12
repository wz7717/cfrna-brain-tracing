"""Append P0-4 sparsity and donor-aware inference tables to the supplement."""

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\Download\cfrna-brain-tracing-streamlit-cloud-ready")
SOURCE = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_complete_updated.docx"
OUTPUT = ROOT / "Bioinformatics_Application_Note_Supplementary_File_P0_complete_with_sparse_donor_tables.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for run in paragraph.runs:
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False, size: float = 8.0) -> None:
    p = cell.paragraphs[0]
    for run in p.runs:
        p._p.remove(run._r)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_table(doc, rows: list[list[str]], widths: list[float]) -> object:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            cell = table.rows[r].cells[c]
            cell.width = Inches(widths[c])
            set_cell(cell, value, bold=(r == 0), size=7.5 if len(widths) == 7 else 8.0)
            if r == 0 or r % 2 == 1:
                shade(cell, "D9E2F3")
    return table


def caption_paragraph(doc, reference_paragraph, text: str):
    paragraph = doc.add_paragraph(style=reference_paragraph.style)
    replace_paragraph_text(paragraph, text)
    return paragraph


def move_before(element, target_paragraph) -> None:
    target_paragraph._p.addprevious(element)


def main() -> None:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    doc = Document(source)

    # Table S5 index receives two entries. Existing Table S5 is document table 7.
    index_table = doc.tables[7]
    for item, content, use in [
        ("Table S12", "Sparse-input sensitivity analysis", "Reports coverage and hierarchical Top3 performance across locked-route sparsity scenarios"),
        ("Table S13", "Donor-aware LOSO/LOMO inference", "Reports donor-macro accuracy, donor bootstrap intervals and cluster-robust comparisons"),
    ]:
        row = index_table.add_row()
        for cell, text in zip(row.cells, [item, content, use]):
            set_cell(cell, text, bold=(cell is row.cells[0]), size=8.0)

    sparse_rows = [
        ["Scenario", "Depth", "Target gene\nretention", "Detected-gene\ncoverage, %", "Network Top3, %\n(bootstrap 95% CI)", "Group Top3, %\n(bootstrap 95% CI)", "Exact Top3, %\n(bootstrap 95% CI)"],
        ["Baseline", "100%", "100%", "99.96", "92.19 (89.60-93.27)", "72.36 (65.51-76.75)", "45.33 (36.81-50.35)"],
        ["Mild", "50%", "80%", "79.17", "91.25 (87.36-92.80)", "71.21 (62.96-75.45)", "41.97 (32.48-47.16)"],
        ["Moderate", "20%", "60%", "58.97", "82.62 (79.25-84.43)", "61.02 (53.97-65.15)", "32.72 (25.86-36.80)"],
        ["Severe", "5%", "40%", "38.58", "71.10 (69.27-72.84)", "50.78 (44.98-54.29)", "25.43 (21.02-28.47)"],
        ["Extreme", "1%", "20%", "18.27", "59.46 (57.40-62.35)", "40.34 (37.65-42.51)", "17.49 (14.00-19.27)"],
    ]
    donor_rows = [
        ["Endpoint", "Paired n\n(donors)", "LOSO donor-macro, %\n(bootstrap 95% CI)", "LOMO donor-macro, %\n(bootstrap 95% CI)", "Robust Δ pp\n(95% CI)", "p raw / BH"],
        ["Network Top3", "819 (9)", "89.69 (86.58-92.41)", "88.51 (85.85-90.98)", "0.98 (0.17-1.78)", "0.0230 / 0.0688"],
        ["Resolution-group Top3", "812 (9)", "68.63 (64.55-73.26)", "67.14 (62.22-72.05)", "3.33 (-0.09-6.74)", "0.0552 / 0.0828"],
        ["Exact-region Top3", "812 (9)", "42.19 (37.21-47.55)", "39.11 (33.09-43.80)", "2.96 (-1.87-7.79)", "0.1959 / 0.1959"],
    ]

    # Add at the document end, then move the new captions and tables before the
    # existing 'Supplementary references' heading, preserving all existing content.
    reference_heading = next(p for p in doc.paragraphs if p.text.strip() == "Supplementary references")
    caption_style_source = next(p for p in doc.paragraphs if p.text.startswith("Table S11."))
    caption_s12 = caption_paragraph(doc, caption_style_source, "Table S12. Sparse-input sensitivity analysis")
    caption_s12.paragraph_format.page_break_before = True
    table_s12 = build_table(doc, sparse_rows, [0.67, 0.55, 0.72, 0.80, 1.35, 1.35, 1.35])
    caption_s13 = caption_paragraph(doc, caption_style_source, "Table S13. Donor-aware LOSO/LOMO inference")
    table_s13 = build_table(doc, donor_rows, [1.10, 0.72, 1.35, 1.35, 1.10, 0.95])
    move_before(caption_s12._p, reference_heading)
    move_before(table_s12._tbl, reference_heading)
    move_before(caption_s13._p, reference_heading)
    move_before(table_s13._tbl, reference_heading)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)
    print("Added Table S12, Table S13, and two Table S5 index rows")


if __name__ == "__main__":
    main()
