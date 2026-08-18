#!/usr/bin/env python
"""Independent, fail-closed verification for the Huang 2025 remediation package."""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


REQUIRED_LEDGER_COLUMNS = [
    "sample_id",
    "fluid",
    "disease_group",
    "tumor_status",
    "patient_id",
    "patient_id_status",
    "expression_available",
    "BrainTrace_output_available",
    "included_in_full159_audit",
    "included_in_CSF_analysis",
    "included_in_plasma_analysis",
    "included_in_tumour_control_analysis",
    "source_QC_status_if_known",
    "source_QC_note",
]
FORBIDDEN_MANUSCRIPT_STRINGS = [
    "minimum paired P",
    "Matched-patient admixture",
    "Each of 77 Huang CSF profiles was paired",
    "77 matched CSF-plasma",
    "59 tumour CSF-plasma pairs",
    "all 77 mixtures",
    "The separate matched CSF-plasma mixture analysis",
    "minimum p=0.304",
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def as_bool(value: str) -> bool:
    return value.strip().lower() in {"true", "1", "yes"}


def document_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_text)


def docx_xml_status(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        word_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        tracked_insertions = len(re.findall(r"<w:ins(?:\s|>)", word_xml))
        tracked_deletions = len(re.findall(r"<w:del(?:\s|>)", word_xml))
        tracked_moves = len(re.findall(r"<w:move(?:From|To)(?:\s|>)", word_xml))
        comments = [name for name in names if "comment" in name.lower()]
        return {
            "tracked_insertions": tracked_insertions,
            "tracked_deletions": tracked_deletions,
            "tracked_moves": tracked_moves,
            "track_revisions_enabled": bool(re.search(r"<w:trackRevisions(?:\s|/|>)", word_xml)),
            "comment_parts": comments,
        }


def require(condition: bool, message: str, errors: list[str]) -> None:
    if not condition:
        errors.append(message)


def verify(args: argparse.Namespace) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    outdir = args.outdir
    expected = {
        "huang_2025_sample_ledger.csv",
        "huang_2025_sample_outputs.csv",
        "huang_2025_network_rankings.csv",
        "huang_2025_exact_region_rankings.csv",
        "huang_2025_network_top1_distribution.csv",
        "huang_2025_network_top3_distribution.csv",
        "huang_2025_fluid_summary.csv",
        "huang_2025_marker_correlations.csv",
        "huang_2025_tumour_control_comparisons.csv",
        "huang_2025_audit_manifest.json",
        "huang_2025_canonical_summary.csv",
        "huang_2025_canonical_summary.json",
        "HUANG_2025_RESULTS.md",
    }
    actual_names = {path.name for path in outdir.iterdir() if path.is_file()}
    require(expected <= actual_names, "Canonical Huang output directory is incomplete.", errors)
    require(
        not any("paired_csf_plasma" in name.lower() or "admixture" in name.lower() for name in actual_names),
        "Canonical Huang output directory contains a retired pseudo-pair artifact.",
        errors,
    )

    summary = json.loads((outdir / "huang_2025_canonical_summary.json").read_text(encoding="utf-8"))
    require(summary.get("protocol_status") == "huang_2025_provenance_remediated", "Unexpected protocol status.", errors)
    require((summary.get("n_profiles"), summary.get("n_csf"), summary.get("n_plasma")) == (159, 77, 82), "Summary cohort counts are not 159/77/82.", errors)
    require(summary.get("n_traceable_outputs") == 159, "Not all 159 outputs are traceable.", errors)
    require(summary.get("patient_paired_analysis") == "NOT_SUPPORTED", "Paired analysis guardrail missing.", errors)
    require(summary.get("synthetic_matched_admixture") == "REMOVED_FROM_CANONICAL_ANALYSIS", "Synthetic-mixture retirement guardrail missing.", errors)

    manifest = json.loads(
        (outdir / "huang_2025_audit_manifest.json").read_text(encoding="utf-8")
    )

    for asset in manifest.get("model_assets", []) + manifest.get("input_assets", []):
        asset_path = str(asset.get("path", ""))
        require(
            bool(asset_path),
            f"Missing provenance path for {asset.get('label')!r}.",
            errors,
        )
        require(
            not Path(asset_path).is_absolute(),
            f"Absolute provenance path found: {asset_path!r}",
            errors,
        )
        require(
            not re.match(r"^[A-Za-z]:[\\/]", asset_path),
            f"Windows absolute provenance path found: {asset_path!r}",
            errors,
        )

    overlap = manifest.get("model_overlap", {})
    require(
        overlap.get("n_source_matrix_features") == 83929,
        "Source feature count is not 83,929.",
        errors,
    )
    require(
        overlap.get("n_selected_input_genes") == 15295,
        "Selected Huang inference-gene count is not 15,295.",
        errors,
    )

    ledger = read_csv(outdir / "huang_2025_sample_ledger.csv")
    require(list(ledger[0]) == REQUIRED_LEDGER_COLUMNS, "Ledger columns differ from the prescribed schema.", errors)
    require(len(ledger) == 159, "Ledger does not contain exactly 159 profiles.", errors)
    require(sum(row["fluid"] == "CSF" for row in ledger) == 77, "Ledger CSF denominator is not 77.", errors)
    require(sum(row["fluid"] == "plasma" for row in ledger) == 82, "Ledger plasma denominator is not 82.", errors)
    require(all(not row["patient_id"].strip() for row in ledger), "Ledger contains an inferred patient identifier.", errors)
    require(all(row["patient_id_status"] == "unknown_not_supplied_in_public_expression_matrix" for row in ledger), "Ledger patient-id status is inconsistent.", errors)
    require(all(as_bool(row["expression_available"]) for row in ledger), "Some ledger profiles lack expression availability.", errors)
    require(all(as_bool(row["BrainTrace_output_available"]) for row in ledger), "Some ledger profiles lack BrainTrace output availability.", errors)

    comparisons = read_csv(outdir / "huang_2025_tumour_control_comparisons.csv")
    require(len(comparisons) == 6, "Expected six tumour-control tests.", errors)
    require(all(row["test"] == "two-sided Mann-Whitney U (profile-level; pairing unavailable)" for row in comparisons), "Unexpected tumour-control test label was found.", errors)
    require({int(row["n_tumour"]) for row in comparisons} == {59, 64}, "Tumour denominators must be 59 and 64.", errors)
    require({int(row["n_control"]) for row in comparisons} == {18}, "Control denominator must be 18.", errors)
    fdrs = [float(row["bh_fdr"]) for row in comparisons]
    require(all(0 <= value <= 1 for value in fdrs), "BH-FDR values fall outside [0,1].", errors)
    require(math.isclose(min(fdrs), float(summary["minimum_bh_fdr"]), rel_tol=0, abs_tol=1e-12), "Summary minimum BH-FDR does not match comparisons.", errors)

    correlations = read_csv(outdir / "huang_2025_marker_correlations.csv")
    require(
        len(correlations) == 32,
        "Expected exactly 32 fluid-by-marker correlation tests.",
        errors,
    )
    require(
        all(row["status"] == "estimated" for row in correlations),
        "A Huang marker correlation is not estimable.",
        errors,
    )
    require(
        all(0 <= float(row["bh_fdr"]) <= 1 for row in correlations),
        "Marker-correlation BH-FDR falls outside [0,1].",
        errors,
    )

    fluid = read_csv(outdir / "huang_2025_fluid_summary.csv")
    for row in fluid:
        n_profiles = int(row["n_profiles"])
        for level in ("top1", "top3"):
            numerator = int(row[f"OMPFC_{level}_numerator"])
            denominator = int(row[f"OMPFC_{level}_denominator"])
            percent = float(row[f"OMPFC_{level}_percent"])
            require(denominator == n_profiles, f"{row['cohort']} {level} denominator differs from cohort size.", errors)
            require(math.isclose(percent, 100 * numerator / denominator, rel_tol=0, abs_tol=1e-9), f"{row['cohort']} {level} percentage is not arithmetic.", errors)

    main_text = document_text(args.main_docx)
    supp_text = document_text(args.supp_docx)
    for forbidden in FORBIDDEN_MANUSCRIPT_STRINGS:
        require(forbidden not in main_text and forbidden not in supp_text, f"Forbidden legacy wording remains: {forbidden!r}", errors)
    required_text = [
        "77 CSF and 82 plasma",
        "No patient-level CSF-plasma correspondence was assumed.",
        "minimum BH-FDR=0.722052",
        "159/159 traceable",
    ]
    merged_text = main_text + "\n" + supp_text
    for phrase in required_text:
        require(phrase in merged_text, f"Required remediation wording is absent: {phrase!r}", errors)

    documents = {}
    for label, path in (("main", args.main_docx), ("supplement", args.supp_docx)):
        status = docx_xml_status(path)
        documents[label] = status
        require(status["tracked_insertions"] == 0, f"{label} DOCX contains tracked insertions.", errors)
        require(status["tracked_deletions"] == 0, f"{label} DOCX contains tracked deletions.", errors)
        require(status["tracked_moves"] == 0, f"{label} DOCX contains tracked moves.", errors)
        require(not status["track_revisions_enabled"], f"{label} DOCX enables Track Changes.", errors)
        require(not status["comment_parts"], f"{label} DOCX contains comment parts.", errors)

    return {
        "status": "PASS" if not errors else "FAIL",
        "errors": errors,
        "warnings": warnings,
        "summary": {
            "n_profiles": summary.get("n_profiles"),
            "n_csf": summary.get("n_csf"),
            "n_plasma": summary.get("n_plasma"),
            "n_traceable_outputs": summary.get("n_traceable_outputs"),
            "minimum_bh_fdr": summary.get("minimum_bh_fdr"),
        },
        "documents": documents,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--outdir", type=Path, default=Path("reproducibility/huang_2025"))
    parser.add_argument("--main-docx", type=Path, required=True)
    parser.add_argument("--supp-docx", type=Path, required=True)
    parser.add_argument("--report", type=Path, required=True)
    args = parser.parse_args()
    result = verify(args)
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
