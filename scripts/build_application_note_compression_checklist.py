from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "manuscript" / "BrainTrace_v5_round4_ApplicationNote限长压缩_修改清单.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(9)
normal.paragraph_format.space_after = Pt(3)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BrainTrace v5 round4：Bioinformatics Application Note 限长压缩修改清单")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(31, 78, 121)

lead = doc.add_paragraph()
lead.add_run("结果：").bold = True
lead.add_run(
    "接受本轮修订后，Word 统计的参考文献前字数由约 6,545 词降至 2,037 词；"
    "全文含参考文献约 2,726 词。目标文章类型保持 Bioinformatics Application Note，"
    "Figure 1 保留为唯一主图。所有删减均为证据分层，不撤回既有计算或限定。"
)

doc.add_heading("一、压缩原则", level=1)
for text in [
    "主稿保留工具定位、三层架构、主要内部/外部验证结果、核心统计假设和使用边界。",
    "审稿新增的完整置信区间、敏感性分析、供体异质性、跨物种解剖学审计、映射偏倚、富集分析和软件细节继续保留在补充材料及计算文件中。",
    "主稿以“结论 + Supplementary 定位”替代成段审稿答复式叙述；未通过缩小字体或页边距规避限长。",
    "本轮对主稿使用真实 Word 修订；补充材料内容无需删除，另存为同轮最新累计版本并完成索引核对。"
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("二、逐部分修改映射", level=1)
rows = [
    ("Abstract", "约263词→精简结构式摘要", "保留工具、三层路线、核心Top3结果、外部审计性质和非临床定位"),
    ("Introduction", "删减同类工具与问题背景展开", "保留任务差异、三层输出和诊断信息"),
    ("System and methods", "约1,849词→约700词级核心方法", "公式、折内防泄漏、F_g非推断性质、跨物种映射和供体聚类推断均保留；细节指向S2-S10"),
    ("Implementation", "压缩工程描述", "保留共享评分核心、速度/内存和一键复现"),
    ("Validation", "约1,989词→约600词级结果摘要", "保留LOSO/LOMO、CRVE区间、层级性能下降、匹配RF、AHBA多标签限制和cfRNA负向域迁移"),
    ("Use and limitations", "约1,050词→约450词级边界声明", "保留Subcortical异质性、视觉通路术语、单区网络退化、cerebellum out-of-scope、3'-degradation与cfRNA限制"),
    ("References/声明", "不以删参考文献换取正文额度", "Funding、COI、贡献、AI披露、数据可用性和必要引用保留")
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ("位置", "压缩动作", "保留内容/证据位置")):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
for row in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("三、关键科学内容保留核对", level=1)
checks = [
    "P0-NEURO1-3：FST/MST、TEO/TE、视觉通路、Subcortical异质性和跨物种限定仍在主稿核心限制中，完整审计在补充材料。",
    "P0-STAT1-3：G=9欠覆盖、独立样本CI仅描述性、CRVE/t8、F_g无F分布推断均保留。",
    "P0-BIO1-3：projected-VSD非数学等价、3'-degradation代理局限、AHBA多标签膨胀均保留。",
    "P1/P2：匹配200基因RF、供体混淆、富集、解剖crosswalk、Friedman/符号翻转假设及软件约束继续由S20和计算文件承载。",
    "主稿不再重复逐轮审稿答复式数字串；所有主要结论均有补充章节或表格定位。"
]
for text in checks:
    doc.add_paragraph("☑ " + text)

doc.add_heading("四、审阅重点", level=1)
doc.add_paragraph(
    "建议审阅者重点确认：(1) 2,037词口径是否符合投稿系统统计；"
    "(2) 主稿中的限定语是否足以支撑Application Note的简洁程度；"
    "(3) Supplementary S2-S20与Tables S11-S17是否覆盖被下沉的证据；"
    "(4) GitHub/Zenodo同步后版本号、URL和Data availability是否需要最终更新。"
)

doc.core_properties.title = "BrainTrace Application Note limit compression checklist"
doc.save(OUT)
print(OUT)
