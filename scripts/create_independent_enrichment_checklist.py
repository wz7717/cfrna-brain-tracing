from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "manuscript" / "BrainTrace_v5_round4_独立GO_KEGG及细胞类型偏倚_修改清单.docx"


doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BrainTrace v5 round 4\n独立 GO/KEGG 及细胞类型偏倚分析修改清单")
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph("日期：2026-07-31　　状态：已累计合并至最新修订稿（保留 Word 修订痕迹）")

doc.add_heading("1. 本轮预先确定的分析层级", level=1)
for text in [
    "主要细胞类型参照：Chiou et al. (2023) 成年恒河猴全脑单细胞多组学图谱。",
    "跨物种敏感性参照：Siletti et al. (2023) 成年人脑转录组细胞图谱。",
    "固定对象：原模型的同一 200-gene Network panel；主要背景为 21,668 个模型空间基因。",
    "GO:BP 与 KEGG 分开控制 FDR；另以数据库注释基因背景进行敏感性分析。",
    "细胞类型采用七个预设家族，单侧超几何检验，并在每个图谱内对七项检验作 BH 校正。",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("2. 计算结果与可审阅结论", level=1)
table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"
headers = ["分析", "家族/数据库", "重叠", "倍数", "校正后 q/FDR", "结论"]
for cell, text in zip(table.rows[0].cells, headers):
    cell.text = text
rows = [
    ("猕猴主分析", "兴奋性神经元", "9/77", "12.66", "2.67×10⁻⁷", "显著"),
    ("猕猴主分析", "抑制性神经元", "6/83", "7.83", "4.20×10⁻⁴", "显著"),
    ("猕猴主分析", "其余五家族", "—", "—", "≥0.207", "均不显著"),
    ("人类敏感性", "兴奋性神经元", "9/407", "2.40", "0.0477", "显著"),
    ("人类敏感性", "抑制性神经元", "13/317", "4.44", "5.89×10⁻⁵", "显著"),
    ("人类敏感性", "其余五家族", "—", "—", "≥0.647", "均不显著"),
    ("GO/KEGG 主背景", "GO:BP / KEGG", "446 / 11 terms", "—", "<0.05", "分别显著"),
    ("GO/KEGG 背景敏感性", "GO:BP / KEGG", "410 / 8 terms", "—", "<0.05", "总体稳定"),
]
for row in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text

doc.add_paragraph(
    "推断边界：两种独立图谱共同支持该固定 panel 的神经元注释偏倚；不能据此断言细胞来源、"
    "病理机制、因果关系或预测有效性。公开 marker 仅为论文提供的 top-ranked genes，且两图谱"
    "粒度不同，因此跨物种结果用于方向性敏感性验证。"
)

doc.add_heading("3. 文稿修改位置", level=1)
items = [
    ("主稿 Validation", "替换原事后注释表述，加入冻结方案、猕猴主分析、人类跨物种敏感性结果及严格推断边界；保持 Application Note 主文限长。"),
    ("补充材料 S20", "补充数据库/API 版本、179/200 标识符映射、背景基因集、校正家族、七类结果、跨物种敏感性、局限及可复现文件索引。"),
    ("补充材料参考文献", "加入 Chiou et al. (2023) 与 Siletti et al. (2023)。"),
    ("图注检查", "主稿及补充图注未发现 GO/KEGG 或细胞来源数值陈述，无需改动；Supplementary Figure S2 仍仅描述 RF 诊断。"),
]
for where, change in items:
    p = doc.add_paragraph()
    p.add_run(where + "：").bold = True
    p.add_run(change)

doc.add_heading("4. 可复现文件", level=1)
for name in [
    "calculations/independent_enrichment/analysis_protocol_20260731.json",
    "calculations/independent_enrichment/independent_celltype_enrichment.csv",
    "calculations/independent_enrichment/independent_primate_marker_sets.csv",
    "calculations/independent_enrichment/independent_human_marker_sets.csv",
    "calculations/independent_enrichment/gprofiler_model_background.csv",
    "calculations/independent_enrichment/gprofiler_annotated_background.csv",
    "calculations/independent_enrichment/gprofiler_GO_BP_representative_components.csv",
    "calculations/independent_enrichment/independent_enrichment_manifest.json",
]:
    doc.add_paragraph(name, style="List Bullet")

doc.add_heading("5. 审阅者核对要点", level=1)
for text in [
    "确认主次表述始终为“独立猕猴图谱主分析；人类图谱跨物种敏感性”。",
    "确认所有新结论均指向 annotation bias，而非 cell of origin 或 mechanism。",
    "确认 200-gene panel、21,668-gene universe 和七项 BH 家族在主稿、补充材料与输出文件中一致。",
    "确认 Word 可显示删除与插入修订，且不存在嵌套修订。",
]:
    doc.add_paragraph("☐ " + text)

doc.core_properties.title = "BrainTrace independent enrichment revision checklist"
doc.core_properties.author = "BrainTrace revision team"
doc.save(OUT)
print(OUT)
